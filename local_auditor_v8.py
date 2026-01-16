#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ACHE SUCATAS DaaS - AUDITOR V8 (OTIMIZADO)
==========================================
Extrai dados de editais usando cascata: JSON → Excel → DOCX → Path → PDF

MELHORIAS V8 (sobre V7.1):
- ✅ CRÍTICO: Extração de URLs de DOCX (88% das URLs estão aqui!)
- ✅ CRÍTICO: Extração de URLs de XLSX
- ✅ CRÍTICO: Processamento de arquivos ZIP (extrai e processa internos)
- ✅ Tags expandidas: DER, Receita Federal, pátio, custódia
- ✅ Keywords de leiloeiro: joaoemilio, leiloesfreire, etc.
- ✅ Regex de datas: desfazimento, alienação

MANTIDO DO V7.1:
- ✅ Cascata de fontes (JSON → Excel → Path → PDF)
- ✅ Correção de encoding UTF-8
- ✅ link_pncp direto do JSON
- ✅ data_leilao com fallback correto

Desenvolvedor: Thiago
Data: 2026-01-15
Versão: 8.0
"""

import re
import pdfplumber
import pandas as pd
import json
import zipfile
import tempfile
import shutil
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Set
import sys

# Tentar importar python-docx (opcional mas recomendado)
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[AVISO] python-docx não instalado. Instale com: pip install python-docx")


# ============================================================================
# CONFIGURAÇÕES
# ============================================================================
PASTA_RAIZ = Path(__file__).parent
PASTA_EDITAIS = PASTA_RAIZ / "ACHE_SUCATAS_DB"
CSV_OUTPUT = PASTA_RAIZ / "analise_editais_v8.csv"

# ============================================================================
# REGEX PATTERNS - V8 EXPANDIDO
# ============================================================================
REGEX_N_EDITAL = re.compile(
    r"(?i)(?:edital|processo|leilão|pregão).*?(\d{1,5}\s*/\s*20\d{2})",
    re.DOTALL
)

REGEX_DATA = re.compile(r"\b(\d{1,2}[/\-]\d{1,2}[/\-]20\d{2})\b")

REGEX_URL = re.compile(
    r"https?://(?:www\.)?[\w\-\.]+(?:/[\w\-\./?%&=]*)?",
    re.IGNORECASE
)

# V8: Regex de datas expandido com mais contextos
REGEX_DATA_CONTEXTUAL = re.compile(
    r"(?i)(?:data|abertura|sessão|leilão|pregão|realização|"
    r"desfazimento|alienação|hasta|arrematação).*?"
    r"(\d{1,2}[/\-]\d{1,2}[/\-]20\d{2})",
    re.DOTALL
)

# ============================================================================
# DICIONÁRIO DE TAGS - V8 EXPANDIDO
# ============================================================================
TAGS_POR_ORGAO = {
    # Originais V7.1
    'prefeitura': 'veiculos_municipais',
    'câmara': 'veiculos_municipais',
    'detran': 'veiculos_detran',
    'polícia': 'veiculos_policiais',
    'tribunal': 'veiculos_judiciarios',
    'tj': 'veiculos_judiciarios',
    
    # NOVOS V8
    'der': 'veiculos_estaduais',
    'departamento de estradas': 'veiculos_estaduais',
    'receita federal': 'bens_apreendidos_federal',
    'secretaria da receita': 'bens_apreendidos_federal',
    'srfb': 'bens_apreendidos_federal',
    'pátio': 'veiculos_detran',
    'patio': 'veiculos_detran',
    'custódia': 'veiculos_detran',
    'custodia': 'veiculos_detran',
    'apreendidos': 'veiculos_detran',
    'pm': 'veiculos_policiais',
    'prf': 'veiculos_policiais',
    'polícia rodoviária': 'veiculos_policiais',
    'policia rodoviaria': 'veiculos_policiais',
    'bombeiro': 'veiculos_bombeiros',
    'samu': 'veiculos_saude',
    'saúde': 'veiculos_saude',
    'saude': 'veiculos_saude',
}

# ============================================================================
# KEYWORDS DE LEILOEIRO - V8 EXPANDIDO
# ============================================================================
KEYWORDS_LEILOEIRO = [
    # Originais V7.1
    'leiloeiro', 'leilao', 'lance', 'arrematacao',
    'superbid', 'sodresantoro', 'zukerman',
    
    # NOVOS V8 (descobertos pelo Comet)
    'joaoemilio',
    'leiloesfreire',
    'frfreiloes',
    'leilaobrasil',
    'leiloes.com',
    'megaleiloes',
    'sold',
    'leilomaster',
    'vipleiloes',
    'leiloesja',
    'portalleiloes',
]


# ============================================================================
# FUNÇÕES AUXILIARES
# ============================================================================

def corrigir_encoding(texto: str) -> str:
    """Corrige dupla codificação UTF-8."""
    if not texto or texto == 'N/D':
        return texto

    try:
        texto_corrigido = texto.encode('latin-1').decode('utf-8')
        return texto_corrigido
    except:
        return texto


def limpar_texto(texto: str, max_length: int = 500) -> str:
    """Limpa e trunca texto."""
    if not texto or texto == 'N/D':
        return texto

    texto = re.sub(r'\n{3,}', '\n\n', texto)
    texto = re.sub(r' {2,}', ' ', texto)

    if len(texto) > max_length:
        texto = texto[:max_length] + '...'

    return texto.strip()


def extrair_data_de_texto(texto: str) -> Optional[str]:
    """Extrai data do leilão de um texto usando contexto."""
    if not texto:
        return None

    # Buscar datas com contexto (V8 expandido)
    matches = REGEX_DATA_CONTEXTUAL.findall(texto)

    if not matches:
        matches = REGEX_DATA.findall(texto)

    for data_str in matches:
        try:
            data_obj = datetime.strptime(data_str.replace('-', '/'), '%d/%m/%Y')
            if data_obj.year >= 2024:
                return data_str
        except:
            continue

    return None


def extrair_urls_de_texto(texto: str) -> List[str]:
    """Extrai todas as URLs de um texto."""
    if not texto:
        return []

    urls = REGEX_URL.findall(texto)
    urls_limpas = [url.rstrip('.,;:)') for url in urls]
    return urls_limpas


def encontrar_link_leiloeiro(urls: List[str]) -> Optional[str]:
    """
    Encontra o link do leiloeiro em uma lista de URLs.
    V8: Keywords expandidas.
    """
    if not urls:
        return None

    # Remover duplicatas mantendo ordem
    urls_unicas = list(dict.fromkeys(urls))

    # Prioridade 1: Keywords específicas de leiloeiro
    for url in urls_unicas:
        url_lower = url.lower()
        if any(kw in url_lower for kw in KEYWORDS_LEILOEIRO):
            return url

    # Prioridade 2: .com.br que não seja gov/pncp
    for url in urls_unicas:
        url_lower = url.lower()
        if '.com' in url_lower and 'gov' not in url_lower and 'pncp' not in url_lower:
            return url

    return None


def encontrar_pasta_dados(pasta_edital: Path) -> Optional[Path]:
    """Encontra a pasta real com os dados."""
    subpastas = [d for d in pasta_edital.iterdir() if d.is_dir()]

    if subpastas:
        return subpastas[0]

    return pasta_edital


# ============================================================================
# EXTRAÇÃO DE DOCX - NOVO V8!
# ============================================================================

def extrair_texto_docx(arquivo_docx: Path) -> str:
    """
    Extrai texto completo de um arquivo DOCX.
    NOVO V8: Essencial para encontrar 88% das URLs!
    """
    if not DOCX_AVAILABLE:
        return ""

    try:
        doc = DocxDocument(str(arquivo_docx))
        texto_completo = []

        # Extrair parágrafos
        for para in doc.paragraphs:
            if para.text.strip():
                texto_completo.append(para.text)

        # Extrair tabelas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texto_completo.append(cell.text)

        return '\n'.join(texto_completo)

    except Exception as e:
        return ""


def extrair_urls_de_docx(pasta_dados: Path) -> List[str]:
    """
    Extrai URLs de todos os arquivos DOCX na pasta.
    NOVO V8!
    """
    urls_encontradas = []

    arquivos_docx = list(pasta_dados.glob("*.docx")) + list(pasta_dados.glob("*.doc"))

    for docx_path in arquivos_docx:
        texto = extrair_texto_docx(docx_path)
        if texto:
            urls = extrair_urls_de_texto(texto)
            urls_encontradas.extend(urls)

    return urls_encontradas


# ============================================================================
# EXTRAÇÃO DE XLSX - NOVO V8!
# ============================================================================

def extrair_urls_de_xlsx(pasta_dados: Path) -> List[str]:
    """
    Extrai URLs de todos os arquivos Excel na pasta.
    NOVO V8: Busca URLs em todas as células!
    """
    urls_encontradas = []

    arquivos_excel = (
        list(pasta_dados.glob("*.xlsx")) +
        list(pasta_dados.glob("*.xls")) +
        list(pasta_dados.glob("*.xlsm"))
    )

    for excel_path in arquivos_excel:
        try:
            # Ler todas as sheets
            xl = pd.ExcelFile(excel_path)
            for sheet_name in xl.sheet_names:
                df = pd.read_excel(xl, sheet_name=sheet_name, dtype=str)

                # Buscar URLs em todas as células
                for col in df.columns:
                    for valor in df[col].dropna():
                        valor_str = str(valor)
                        if 'http' in valor_str.lower():
                            urls = extrair_urls_de_texto(valor_str)
                            urls_encontradas.extend(urls)

        except Exception as e:
            pass

    return urls_encontradas


# ============================================================================
# PROCESSAMENTO DE ZIP - NOVO V8!
# ============================================================================

def extrair_urls_de_zip(pasta_dados: Path) -> List[str]:
    """
    Extrai URLs de arquivos dentro de ZIPs.
    NOVO V8: Processa PDF/DOCX/XLSX dentro de arquivos compactados!
    """
    urls_encontradas = []

    arquivos_zip = list(pasta_dados.glob("*.zip"))

    for zip_path in arquivos_zip:
        try:
            # Criar diretório temporário
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir)

                # Extrair ZIP
                with zipfile.ZipFile(zip_path, 'r') as zf:
                    zf.extractall(temp_path)

                # Processar arquivos extraídos
                # DOCX
                urls_encontradas.extend(extrair_urls_de_docx(temp_path))

                # XLSX
                urls_encontradas.extend(extrair_urls_de_xlsx(temp_path))

                # PDF
                for pdf_path in temp_path.glob("**/*.pdf"):
                    try:
                        with pdfplumber.open(pdf_path) as pdf:
                            for page in pdf.pages:
                                texto = page.extract_text() or ""
                                urls = extrair_urls_de_texto(texto)
                                urls_encontradas.extend(urls)
                    except:
                        pass

        except Exception as e:
            pass

    return urls_encontradas


# ============================================================================
# FUNÇÕES DE EXTRAÇÃO - CASCATA DE FONTES V8
# ============================================================================

def extrair_de_metadados_json(pasta_dados: Path) -> Dict[str, str]:
    """
    PRIORIDADE 1: Extrai dados do arquivo metadados_pncp.json
    """
    dados = {}
    json_file = pasta_dados / "metadados_pncp.json"

    if not json_file.exists():
        return dados

    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            metadados = json.load(f)

        # CAMPO: titulo
        if 'titulo' in metadados and metadados['titulo']:
            titulo_original = metadados['titulo']
            titulo_corrigido = corrigir_encoding(titulo_original)
            dados['titulo'] = titulo_corrigido

            match = REGEX_N_EDITAL.search(titulo_corrigido)
            if match:
                dados['n_edital'] = match.group(1).replace(' ', '')

        # CAMPO: descricao
        descricao_original = ''

        if 'descricao' in metadados and metadados['descricao']:
            descricao_original = metadados['descricao']
            descricao_corrigida = corrigir_encoding(descricao_original)
            descricao_limpa = limpar_texto(descricao_corrigida, max_length=500)
            dados['descricao'] = descricao_limpa
        elif 'objeto' in metadados and metadados['objeto']:
            descricao_original = metadados['objeto']
            objeto_corrigido = corrigir_encoding(descricao_original)
            dados['descricao'] = limpar_texto(objeto_corrigido, max_length=500)

        # CAMPO: data_leilao
        if 'data_inicio_propostas' in metadados and metadados['data_inicio_propostas']:
            dados['data_leilao'] = metadados['data_inicio_propostas'].split('T')[0]
        elif 'dataAberturaPropostas' in metadados and metadados['dataAberturaPropostas']:
            dados['data_leilao'] = metadados['dataAberturaPropostas'].split('T')[0]
        elif descricao_original:
            data_extraida = extrair_data_de_texto(descricao_original)
            if data_extraida:
                dados['data_leilao'] = data_extraida

        # CAMPO: link_leiloeiro - Busca em múltiplos campos do JSON
        urls_todas: List[str] = []

        if descricao_original:
            urls_todas.extend(extrair_urls_de_texto(descricao_original))

        if 'objeto' in metadados and metadados['objeto']:
            urls_todas.extend(extrair_urls_de_texto(metadados['objeto']))

        if 'informacoes_complementares' in metadados:
            urls_todas.extend(extrair_urls_de_texto(str(metadados['informacoes_complementares'])))

        if 'files_meta' in metadados:
            for file_info in metadados.get('files_meta', []):
                if 'titulo' in file_info:
                    urls_todas.extend(extrair_urls_de_texto(file_info['titulo']))

        link_leiloeiro = encontrar_link_leiloeiro(urls_todas)
        if link_leiloeiro:
            dados['link_leiloeiro'] = link_leiloeiro

        # CAMPO: link_pncp
        if 'link_pncp' in metadados and metadados['link_pncp']:
            dados['link_pncp'] = metadados['link_pncp']

        # CAMPO: orgao
        if 'orgao_nome' in metadados and metadados['orgao_nome']:
            dados['orgao'] = metadados['orgao_nome']

        # CAMPO: uf
        if 'uf' in metadados and metadados['uf']:
            dados['uf'] = metadados['uf']

        # CAMPO: municipio/cidade
        if 'municipio' in metadados and metadados['municipio']:
            dados['cidade'] = metadados['municipio']

    except Exception as e:
        print(f"[AVISO] Erro ao ler JSON: {e}")

    return dados


def extrair_de_excel(pasta_dados: Path) -> Dict[str, str]:
    """PRIORIDADE 2: Extrai dados de arquivos Excel."""
    dados = {}

    arquivos_excel = (
        list(pasta_dados.glob("*.xlsx")) +
        list(pasta_dados.glob("*.xls")) +
        list(pasta_dados.glob("*.csv"))
    )

    if not arquivos_excel:
        return dados

    for excel_path in arquivos_excel:
        try:
            if excel_path.suffix == '.csv':
                df = pd.read_csv(excel_path, encoding='utf-8', nrows=50)
            else:
                df = pd.read_excel(excel_path, nrows=50)

            if df.empty:
                continue

            for col in df.columns:
                col_lower = str(col).lower()

                if 'edital' in col_lower and 'n_edital' not in dados:
                    valores = df[col].dropna().astype(str).tolist()
                    if valores:
                        dados['n_edital'] = valores[0]

                if any(kw in col_lower for kw in ['data', 'leilao', 'abertura']):
                    if 'data_leilao_excel' not in dados:
                        valores = df[col].dropna().astype(str).tolist()
                        if valores:
                            data = extrair_data_de_texto(valores[0])
                            if data:
                                dados['data_leilao_excel'] = data

                if any(kw in col_lower for kw in ['desc', 'objeto', 'bem']):
                    if 'descricao_excel' not in dados:
                        valores = df[col].dropna().astype(str).tolist()
                        if valores:
                            dados['descricao_excel'] = ' | '.join(valores[:3])

        except Exception as e:
            pass

    return dados


def extrair_de_docx(pasta_dados: Path) -> Dict[str, str]:
    """
    PRIORIDADE 2.5: Extrai dados de arquivos DOCX.
    NOVO V8!
    """
    dados = {}

    if not DOCX_AVAILABLE:
        return dados

    arquivos_docx = list(pasta_dados.glob("*.docx")) + list(pasta_dados.glob("*.doc"))

    for docx_path in arquivos_docx:
        texto = extrair_texto_docx(docx_path)
        if not texto:
            continue

        # Extrair n_edital
        if 'n_edital' not in dados:
            match = REGEX_N_EDITAL.search(texto)
            if match:
                dados['n_edital'] = match.group(1).replace(' ', '')

        # Extrair data_leilao
        if 'data_leilao' not in dados:
            data = extrair_data_de_texto(texto)
            if data:
                dados['data_leilao'] = data

        # Extrair URLs para link_leiloeiro
        if 'link_leiloeiro' not in dados:
            urls = extrair_urls_de_texto(texto)
            link = encontrar_link_leiloeiro(urls)
            if link:
                dados['link_leiloeiro'] = link

    return dados


def extrair_de_path(pasta_edital: Path) -> Dict[str, str]:
    """PRIORIDADE 3: Extrai dados do nome da pasta PRINCIPAL."""
    dados = {}

    partes = pasta_edital.name.split('_')

    if len(partes) >= 3:
        dados['uf'] = partes[0].upper()
        dados['cidade'] = partes[1].replace('-', ' ').title()
        dados['orgao'] = ' '.join(partes[2:]).replace('-', ' ').title()
    elif len(partes) >= 2:
        dados['uf'] = partes[0].upper()
        dados['cidade'] = partes[1].replace('-', ' ').title()

    return dados


def extrair_de_pdf(pasta_dados: Path) -> Dict[str, str]:
    """PRIORIDADE 4 (FALLBACK): Extrai dados do PDF."""
    dados = {}

    arquivos_pdf = list(pasta_dados.glob("*.pdf"))
    if not arquivos_pdf:
        return dados

    pdf_path = arquivos_pdf[0]

    try:
        with pdfplumber.open(pdf_path) as pdf:
            texto_completo = ""
            for page in pdf.pages:
                texto_completo += page.extract_text() or ""

            if not texto_completo.strip():
                return dados

            # Extrair n_edital
            match = REGEX_N_EDITAL.search(texto_completo)
            if match:
                dados['n_edital'] = match.group(1).replace(' ', '')

            # Extrair data do leilão
            data_extraida = extrair_data_de_texto(texto_completo)
            if data_extraida:
                dados['data_leilao'] = data_extraida

            # Extrair título contextual
            match_titulo = re.search(
                r"(?i)(?:objeto|finalidade).*?[:;]\s*(.{10,200}?)(?:\.|;|\n\n)",
                texto_completo,
                re.DOTALL
            )
            if match_titulo:
                dados['titulo'] = match_titulo.group(1).strip()[:200]

            # Extrair link do leiloeiro
            urls = extrair_urls_de_texto(texto_completo)
            link_leiloeiro = encontrar_link_leiloeiro(urls)
            if link_leiloeiro:
                dados['link_leiloeiro'] = link_leiloeiro

    except Exception as e:
        print(f"[AVISO] Erro ao processar PDF: {e}")

    return dados


def extrair_urls_de_todos_arquivos(pasta_dados: Path) -> List[str]:
    """
    V8: Extrai URLs de TODOS os tipos de arquivo.
    Combina DOCX + XLSX + ZIP para maximizar cobertura.
    """
    urls_todas: List[str] = []

    # DOCX (88% das URLs!)
    urls_todas.extend(extrair_urls_de_docx(pasta_dados))

    # XLSX
    urls_todas.extend(extrair_urls_de_xlsx(pasta_dados))

    # ZIP (pode conter DOCX/XLSX/PDF)
    urls_todas.extend(extrair_urls_de_zip(pasta_dados))

    return urls_todas


def gerar_tags(orgao: str, uf: str) -> str:
    """
    Gera tags baseadas no órgão.
    V8: Dicionário expandido.
    """
    if not orgao or orgao == 'N/D':
        return 'veiculos_gerais'

    orgao_lower = orgao.lower()

    for palavra_chave, tag in TAGS_POR_ORGAO.items():
        if palavra_chave in orgao_lower:
            return tag

    return 'veiculos_gerais'


# ============================================================================
# PROCESSAMENTO PRINCIPAL
# ============================================================================

def processar_edital(pasta_edital: Path) -> Dict[str, str]:
    """Processa um edital usando CASCATA DE FONTES V8."""
    print(f"[INFO] Processando: {pasta_edital.name}")

    pasta_dados = encontrar_pasta_dados(pasta_edital)

    if not pasta_dados:
        print(f"  [ERRO] Pasta de dados não encontrada")
        return None

    if pasta_dados != pasta_edital:
        print(f"  [INFO] Pasta de dados: {pasta_dados.name}")

    # Inicializar dados
    dados_finais = {
        'n_edital': 'N/D',
        'data_leilao': 'N/D',
        'titulo': 'N/D',
        'descricao': 'N/D',
        'orgao': 'N/D',
        'uf': 'N/D',
        'cidade': 'N/D',
        'tags': 'N/D',
        'link_leiloeiro': 'N/D',
        'link_pncp': 'N/D',
        'arquivo_origem': pasta_edital.name
    }

    # CASCATA DE FONTES V8 (inclui DOCX!)
    fontes = [
        ('JSON', extrair_de_metadados_json(pasta_dados)),
        ('Excel', extrair_de_excel(pasta_dados)),
        ('DOCX', extrair_de_docx(pasta_dados)),  # NOVO V8!
        ('Path', extrair_de_path(pasta_edital)),
        ('PDF', extrair_de_pdf(pasta_dados))
    ]

    # Mesclar dados
    campos_encontrados = []
    for nome_fonte, dados_fonte in fontes:
        for campo, valor in dados_fonte.items():
            if campo in dados_finais and (dados_finais[campo] == 'N/D' or not dados_finais[campo]):
                if valor and valor != 'N/D':
                    dados_finais[campo] = valor
                    campos_encontrados.append(f"{campo}({nome_fonte})")

    # V8: Busca INTENSIVA de link_leiloeiro em TODOS os arquivos
    if dados_finais['link_leiloeiro'] == 'N/D':
        urls_todos_arquivos = extrair_urls_de_todos_arquivos(pasta_dados)
        link_encontrado = encontrar_link_leiloeiro(urls_todos_arquivos)
        if link_encontrado:
            dados_finais['link_leiloeiro'] = link_encontrado
            campos_encontrados.append("link_leiloeiro(V8_DEEP)")

    if campos_encontrados:
        print(f"  [OK] Extraído: {', '.join(campos_encontrados[:5])}{'...' if len(campos_encontrados) > 5 else ''}")
    else:
        print(f"  [AVISO] Nenhum dado extraído")

    # Gerar tags (V8 expandido)
    dados_finais['tags'] = gerar_tags(dados_finais['orgao'], dados_finais['uf'])

    return dados_finais


def varrer_editais() -> List[Dict[str, str]]:
    """Varre todos os editais."""
    if not PASTA_EDITAIS.exists():
        print(f"[ERRO] Diretório {PASTA_EDITAIS} não encontrado")
        sys.exit(1)

    print(f"[INFO] Varrendo editais em: {PASTA_EDITAIS}")

    resultados = []
    pastas_editais = [d for d in PASTA_EDITAIS.iterdir() if d.is_dir()]

    print(f"[INFO] Total de pastas encontradas: {len(pastas_editais)}\n")

    for i, pasta in enumerate(pastas_editais, 1):
        try:
            print(f"[{i}/{len(pastas_editais)}] ", end='')
            dados = processar_edital(pasta)
            if dados:
                resultados.append(dados)
        except Exception as e:
            print(f"[ERRO] Falha ao processar {pasta.name}: {e}")

    return resultados


def salvar_csv(dados: List[Dict[str, str]]):
    """Salva resultados em CSV."""
    if not dados:
        print("[AVISO] Nenhum dado para salvar")
        return

    df = pd.DataFrame(dados)
    df.to_csv(CSV_OUTPUT, index=False, encoding='utf-8-sig')

    print(f"\n[OK] Arquivo gerado: {CSV_OUTPUT}")
    print(f"[OK] Total de editais processados: {len(dados)}")

    # Estatísticas de preenchimento
    print("\n" + "=" * 60)
    print("ESTATÍSTICAS DE PREENCHIMENTO - V8")
    print("=" * 60)

    for col in df.columns:
        if col == 'arquivo_origem':
            continue

        total = len(df)
        nao_vazio = df[col].notna() & (df[col] != '') & (df[col] != 'N/D')
        preenchidos = nao_vazio.sum()
        taxa = (preenchidos / total * 100) if total > 0 else 0

        status = "✓" if taxa == 100 else "⚠" if taxa >= 50 else "✗"
        barra = "█" * int(taxa / 5)
        print(f"{status} {col:20s} {preenchidos:3d}/{total:3d} ({taxa:5.1f}%) {barra}")


# ============================================================================
# MAIN
# ============================================================================

def main():
    print("\n" + "=" * 60)
    print("ACHE SUCATAS DaaS - AUDITOR V8 (OTIMIZADO)")
    print("=" * 60)
    print("\nMELHORIAS V8:")
    print("✅ CRÍTICO: Extração de URLs de DOCX (88% das URLs!)")
    print("✅ CRÍTICO: Extração de URLs de XLSX")
    print("✅ CRÍTICO: Processamento de arquivos ZIP")
    print("✅ Tags expandidas: DER, Receita Federal, pátio, custódia")
    print("✅ Keywords leiloeiro: joaoemilio, leiloesfreire, etc.")
    print("✅ Regex datas: desfazimento, alienação")
    print("\nMANTIDO DO V7.1:")
    print("✅ Cascata: JSON → Excel → DOCX → Path → PDF")
    print("✅ Encoding UTF-8 corrigido")
    print("✅ link_pncp direto do JSON")

    if not DOCX_AVAILABLE:
        print("\n⚠️  AVISO: python-docx não instalado!")
        print("   Instale com: pip install python-docx")
        print("   Sem isso, extração de DOCX não funcionará.")

    print("=" * 60 + "\n")

    # Processar todos os editais
    resultados = varrer_editais()

    # Salvar CSV
    salvar_csv(resultados)

    print("\n" + "=" * 60)
    print("PROCESSAMENTO CONCLUÍDO - V8")
    print("=" * 60 + "\n")


if __name__ == "__main__":
    main()
