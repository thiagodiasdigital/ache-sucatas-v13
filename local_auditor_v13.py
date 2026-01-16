#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
ACHE SUCATAS DaaS - AUDITOR V13 - INTEGRAÇÃO SUPABASE
Extrai dados de editais usando cascata: JSON -> Excel -> DOCX -> Path -> PDF
Persiste dados no Supabase PostgreSQL + backup CSV/XLSX

V13 (2026-01-16):
- NOVO: Integração com Supabase PostgreSQL
- NOVO: RLS (Row Level Security) ativado
- NOVO: Dual storage (Supabase + CSV/XLSX backup)
- NOVO: SupabaseRepository com error handling robusto
- NOVO: Logs seguros (sem dados sensíveis)
- MANTÉM: Todas as features do V12 (API PNCP, cascata, etc.)

V12:
- API PNCP integrada como FONTE 0 (data_leilao + valor_estimado)
- Correções críticas de bugs (links, datas, tags, títulos)
- 100% cobertura em data_leilao e valor_estimado

V11:
- FIX: Adiciona .leilao.br ao SUPPORTED_TLDS
- FIX: Captura datas extensas (ex: "14 de janeiro de 2026")
- FIX: Corrige formato link_pncp
- NOVO: Adiciona id_interno, data_publicacao, data_atualizacao, objeto_resumido
- NOVO: Gera arquivo RESULTADO_FINAL.xlsx com todos os campos
"""

from __future__ import annotations

import hashlib
import json
import re
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import pandas as pd
import pdfplumber

# V13: Importar SupabaseRepository
from supabase_repository import SupabaseRepository

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[AVISO] python-docx nao instalado. Instale com: pip install python-docx")

# V12: Lista de domínios PROIBIDOS (não são sites de leiloeiros)
DOMINIOS_INVALIDOS = {
    'hotmail.com', 'hotmail.com.br',
    'yahoo.com', 'yahoo.com.br',
    'gmail.com', 'outlook.com',
    'uol.com.br', 'bol.com.br',
    'terra.com.br', 'ig.com.br',
    'globo.com', 'msn.com',
    'live.com', 'icloud.com'
}

# V12: Dicionário de tags específicas
MAPA_TAGS = {
    'sucata': ['sucata', 'sucateamento'],
    'documentado': ['documentado', 'com documento'],
    'sem_documento': ['sem documento', 'indocumentado'],
    'sinistrado': ['sinistrado', 'acidentado'],
    'automovel': ['automóvel', 'automovel', 'carro'],
    'motocicleta': ['motocicleta', 'moto'],
    'caminhao': ['caminhão', 'caminhao'],
    'onibus': ['ônibus', 'onibus'],
    'utilitario': ['utilitário', 'pick-up', 'van'],
    'apreendido': ['apreendido', 'apreensão']
}


PASTA_RAIZ = Path(__file__).parent
PASTA_EDITAIS = PASTA_RAIZ / "ACHE_SUCATAS_DB"
CSV_OUTPUT = PASTA_RAIZ / "analise_editais_v13.csv"  # V13: novo nome


REGEX_N_EDITAL = re.compile(
    r"(?i)(?:edital|processo|leilao|pregao).*?(\d{1,5}\s*/\s*20\d{2})",
    re.DOTALL,
)

REGEX_DATA = re.compile(r"\b(\d{1,2}[/\-]\d{1,2}[/\-]20\d{2})\b")

REGEX_DATA_CONTEXTUAL = re.compile(
    r"(?i)(?:data|abertura|sessao|leilao|pregao|realizacao|"
    r"desfazimento|alienacao|hasta|arrematacao).*?"
    r"(\d{1,2}[/\-]\d{1,2}[/\-]20\d{2})",
    re.DOTALL,
)

# V11: Suporte para datas extensas como "14 de janeiro de 2026"
MESES_BR = {
    "janeiro": "01", "fevereiro": "02", "março": "03", "marco": "03",
    "abril": "04", "maio": "05", "junho": "06", "julho": "07",
    "agosto": "08", "setembro": "09", "outubro": "10",
    "novembro": "11", "dezembro": "12"
}

REGEX_DATA_EXTENSO = re.compile(
    r"(\d{1,2})\s*de\s*(janeiro|fevereiro|março|marco|abril|maio|junho|"
    r"julho|agosto|setembro|outubro|novembro|dezembro)\s*de\s*(20\d{2})",
    re.IGNORECASE
)

# V11: Regex para extrair resumo dos objetos/veículos do leilão
REGEX_OBJETO_RESUMIDO = re.compile(
    r"(?i)(?:lote|item|veiculo|veículo|placa|marca|modelo|chassi).*?[:;]\s*(.{20,300})",
    re.DOTALL
)

SUPPORTED_TLDS = r"(?:\.leilao\.br|\.com\.br|\.com|\.org\.br|\.gov\.br|\.net\.br|\.net)"

REGEX_URL = re.compile(
    rf"(?:https?://)?(?:www\.)?[\w\-\.]+{SUPPORTED_TLDS}"
    r"(?:/[\w\-\./?%&=~#]*)?",
    re.IGNORECASE,
)

REGEX_URL_QUEBRADA = re.compile(
    rf"(?:https?://)?(?:www\.)?[\w\-\.]+\s*[\n\r]\s*{SUPPORTED_TLDS}"
    r"(?:/[\w\-\./?%&=~#]*)?",
    re.IGNORECASE,
)

REGEX_URL_ENCURTADA = re.compile(
    r"(?:bit\.ly|goo\.gl|tinyurl\.com|t\.co)/[\w\-]+",
    re.IGNORECASE,
)

REGEX_PLATAFORMA_CONTEXTUAL = re.compile(
    rf"(?i)(?:plataforma|acesse|site|portal|sistema|endereco\s*eletronico|"
    rf"disponivel\s*em|link|url)[:\s]*"
    rf"((?:https?://)?(?:www\.)?[\w\-\.]+{SUPPORTED_TLDS}"
    r"(?:/[\w\-\./?%&=~#]*)?)",
    re.IGNORECASE,
)

TAGS_POR_ORGAO = {
    "prefeitura": "veiculos_municipais",
    "camara": "veiculos_municipais",
    "detran": "veiculos_detran",
    "policia": "veiculos_policiais",
    "tribunal": "veiculos_judiciarios",
    "tj": "veiculos_judiciarios",
    "der": "veiculos_estaduais",
    "departamento de estradas": "veiculos_estaduais",
    "receita federal": "bens_apreendidos_federal",
    "secretaria da receita": "bens_apreendidos_federal",
    "srfb": "bens_apreendidos_federal",
    "patio": "veiculos_detran",
    "custodia": "veiculos_detran",
    "apreendidos": "veiculos_detran",
    "pm": "veiculos_policiais",
    "prf": "veiculos_policiais",
    "policia rodoviaria": "veiculos_policiais",
    "bombeiro": "veiculos_bombeiros",
    "samu": "veiculos_saude",
    "saude": "veiculos_saude",
}

KEYWORDS_LEILOEIRO = [
    "leiloeiro",
    "leilao",
    "lance",
    "arrematacao",
    "superbid",
    "sodresantoro",
    "zukerman",
    "joaoemilio",
    "leiloesfreire",
    "frfreiloes",
    "leilaobrasil",
    "megaleiloes",
    "sold",
    "leilomaster",
    "vipleiloes",
    "leiloesja",
    "portalleiloes",
    "lanceja",
    "leiloar",
    "leiloado",
    "hastaleiloes",
    "cleiloes",
    "bfreiloes",
    "alfreidoleiloes",
    "rioleiloes",
    "leilonet",
    "leilocentro",
]


@dataclass(frozen=True)
class ResultadoEdital:
    id_interno: str = "N/D"            # V11: PK interna
    n_edital: str = "N/D"
    data_publicacao: str = "N/D"       # V11: do JSON
    data_atualizacao: str = "N/D"      # V11: do JSON
    data_leilao: str = "N/D"
    titulo: str = "N/D"
    descricao: str = "N/D"
    objeto_resumido: str = "N/D"       # V11: extrair do PDF
    orgao: str = "N/D"
    uf: str = "N/D"
    cidade: str = "N/D"
    tags: str = "N/D"
    link_leiloeiro: str = "N/D"
    link_pncp: str = "N/D"
    modalidade_leilao: str = "N/D"
    valor_estimado: str = "N/D"
    quantidade_itens: str = "N/D"
    nome_leiloeiro: str = "N/D"
    arquivo_origem: str = "N/D"

    def as_dict(self) -> Dict[str, str]:
        return {
            "id_interno": self.id_interno,
            "n_edital": self.n_edital,
            "data_publicacao": self.data_publicacao,
            "data_atualizacao": self.data_atualizacao,
            "data_leilao": self.data_leilao,
            "titulo": self.titulo,
            "descricao": self.descricao,
            "objeto_resumido": self.objeto_resumido,
            "orgao": self.orgao,
            "uf": self.uf,
            "cidade": self.cidade,
            "tags": self.tags,
            "link_leiloeiro": self.link_leiloeiro,
            "link_pncp": self.link_pncp,
            "modalidade_leilao": self.modalidade_leilao,
            "valor_estimado": self.valor_estimado,
            "quantidade_itens": self.quantidade_itens,
            "nome_leiloeiro": self.nome_leiloeiro,
            "arquivo_origem": self.arquivo_origem,
        }


def corrigir_encoding(texto: str) -> str:
    if not texto or texto == "N/D":
        return texto
    try:
        return texto.encode("latin-1").decode("utf-8")
    except (UnicodeDecodeError, UnicodeEncodeError):
        return texto


def limpar_texto(texto: str, max_length: int = 500) -> str:
    if not texto or texto == "N/D":
        return texto
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    texto = re.sub(r" {2,}", " ", texto)
    if len(texto) > max_length:
        texto = texto[:max_length] + "..."
    return texto.strip()


def converter_data_extenso(match: re.Match) -> str:
    """Converte '14 de janeiro de 2026' para '14/01/2026'"""
    dia = match.group(1).zfill(2)
    mes = MESES_BR.get(match.group(2).lower(), "01")
    ano = match.group(3)
    return f"{dia}/{mes}/{ano}"


def gerar_id_interno(link_pncp: str) -> str:
    """Gera ID interno único baseado no link_pncp ou timestamp"""
    if not link_pncp or link_pncp == "N/D":
        return f"ID_{datetime.now().strftime('%Y%m%d%H%M%S%f')}"
    return f"ID_{hashlib.md5(link_pncp.encode()).hexdigest()[:12].upper()}"


def extrair_data_de_texto(texto: str) -> Optional[str]:
    """Extrai data do texto, incluindo formatos extensos como '14 de janeiro de 2026'"""
    if not texto:
        return None

    # Tentativa 1: Regex contextual (dd/mm/yyyy)
    matches = REGEX_DATA_CONTEXTUAL.findall(texto)
    if not matches:
        matches = REGEX_DATA.findall(texto)

    for data_str in matches:
        try:
            data_obj = datetime.strptime(data_str.replace("-", "/"), "%d/%m/%Y")
            if data_obj.year >= 2024:
                return data_str
        except ValueError:
            continue

    # Tentativa 2: Formato extenso "14 de janeiro de 2026"
    match_extenso = REGEX_DATA_EXTENSO.search(texto)
    if match_extenso:
        data_convertida = converter_data_extenso(match_extenso)
        try:
            data_obj = datetime.strptime(data_convertida, "%d/%m/%Y")
            if data_obj.year >= 2024:
                return data_convertida
        except ValueError:
            pass

    return None


def extrair_objeto_resumido(texto: str) -> str:
    """Extrai resumo dos objetos/veículos do leilão"""
    if not texto:
        return "N/D"

    matches = REGEX_OBJETO_RESUMIDO.findall(texto)
    if matches:
        # Limpa e junta até 3 matches
        objetos = [m.strip()[:300] for m in matches[:3]]
        return " | ".join(objetos)[:300]

    return "N/D"


def normalizar_url(url: str) -> str:
    url = url.strip()
    url = re.sub(r"\s+", "", url)
    url = url.rstrip(".,;:)>")
    url = url.rstrip("\"'")
    if not url:
        return ""
    if not url.lower().startswith(("http://", "https://")):
        url = "https://" + url
    return url


def extrair_urls_de_texto(texto: str) -> List[str]:
    if not texto:
        return []

    urls: List[str] = []

    urls.extend(REGEX_URL.findall(texto))

    for raw in REGEX_URL_QUEBRADA.findall(texto):
        urls.append(raw)

    for raw in REGEX_URL_ENCURTADA.findall(texto):
        urls.append(raw)

    urls_limpas: List[str] = []
    for raw in urls:
        url = normalizar_url(raw)
        if url and len(url) > 10:
            urls_limpas.append(url)

    return list(dict.fromkeys(urls_limpas))


def extrair_urls_contextuais(texto: str) -> List[str]:
    if not texto:
        return []

    urls: List[str] = []
    for match in REGEX_PLATAFORMA_CONTEXTUAL.findall(texto):
        url = normalizar_url(match)
        if url and len(url) > 10:
            urls.append(url)

    return list(dict.fromkeys(urls))


def encontrar_link_leiloeiro(urls: Sequence[str], texto_completo: str = "") -> Optional[str]:
    urls_unicas = list(dict.fromkeys([u for u in urls if u]))

    for url in urls_unicas:
        url_lower = url.lower()
        if any(keyword in url_lower for keyword in KEYWORDS_LEILOEIRO):
            return url

    if texto_completo:
        urls_ctx = extrair_urls_contextuais(texto_completo)
        for url in urls_ctx:
            url_lower = url.lower()
            if "gov" in url_lower or "pncp" in url_lower:
                continue
            if any(keyword in url_lower for keyword in KEYWORDS_LEILOEIRO):
                return url

        for url in urls_ctx:
            url_lower = url.lower()
            if "gov" in url_lower or "pncp" in url_lower:
                continue
            if ".com.br" in url_lower or ".com" in url_lower or ".net.br" in url_lower or ".net" in url_lower:
                return url

    for url in urls_unicas:
        url_lower = url.lower()
        if "gov" in url_lower or "pncp" in url_lower:
            continue
        if any(tld in url_lower for tld in [".com.br", ".com", ".net.br", ".net", ".org.br"]):
            return url

    return None


def gerar_tags(orgao: str) -> str:
    if not orgao or orgao == "N/D":
        return "veiculos_gerais"

    orgao_lower = orgao.lower()
    for palavra_chave, tag in TAGS_POR_ORGAO.items():
        if palavra_chave in orgao_lower:
            return tag

    return "veiculos_gerais"


def extrair_texto_docx(arquivo_docx: Path) -> str:
    if not DOCX_AVAILABLE:
        return ""
    try:
        doc = DocxDocument(str(arquivo_docx))
        partes: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                partes.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        partes.append(cell.text)
        return "\n".join(partes)
    except Exception:
        return ""


def extrair_urls_de_docx(pasta: Path) -> List[str]:
    urls: List[str] = []
    for docx_path in list(pasta.glob("*.docx")) + list(pasta.glob("*.doc")):
        texto = extrair_texto_docx(docx_path)
        if not texto:
            continue
        urls.extend(extrair_urls_de_texto(texto))
        urls.extend(extrair_urls_contextuais(texto))
    return list(dict.fromkeys(urls))


def extrair_urls_de_xlsx(pasta: Path) -> List[str]:
    urls: List[str] = []
    arquivos = list(pasta.glob("*.xlsx")) + list(pasta.glob("*.xls")) + list(pasta.glob("*.xlsm"))
    for excel_path in arquivos:
        try:
            xl = pd.ExcelFile(excel_path)
            for sheet_name in xl.sheet_names:
                df = pd.read_excel(xl, sheet_name=sheet_name, dtype=str)
                for col in df.columns:
                    for valor in df[col].dropna():
                        valor_str = str(valor)
                        if "http" in valor_str.lower() or ".com" in valor_str.lower() or ".net" in valor_str.lower():
                            urls.extend(extrair_urls_de_texto(valor_str))
        except Exception:
            continue
    return list(dict.fromkeys(urls))


def extrair_urls_de_zip(pasta: Path) -> List[str]:
    urls: List[str] = []
    for zip_path in pasta.glob("*.zip"):
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir)
                with zipfile.ZipFile(zip_path, "r") as zf:
                    zf.extractall(temp_path)

                urls.extend(extrair_urls_de_docx(temp_path))
                urls.extend(extrair_urls_de_xlsx(temp_path))

                for pdf_path in temp_path.glob("**/*.pdf"):
                    try:
                        texto = extrair_texto_pdfs([pdf_path])
                        urls.extend(extrair_urls_de_texto(texto))
                        urls.extend(extrair_urls_contextuais(texto))
                    except Exception:
                        continue
        except Exception:
            continue
    return list(dict.fromkeys(urls))


def extrair_texto_pdfs(pdfs: Sequence[Path]) -> str:
    partes: List[str] = []
    for pdf_path in pdfs:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    partes.append(page.extract_text() or "")
        except Exception:
            continue
    return " ".join([p for p in partes if p]).strip()


def extrair_de_metadados_json(pasta: Path) -> Dict[str, str]:
    dados: Dict[str, str] = {}
    json_file = pasta / "metadados_pncp.json"
    if not json_file.exists():
        return dados

    try:
        with open(json_file, "r", encoding="utf-8") as f:
            metadados = json.load(f)

        # V12: Guardar JSON raw para funções V12
        dados["_json_raw"] = metadados

        titulo_original = metadados.get("titulo") or ""
        if titulo_original:
            titulo_corrigido = corrigir_encoding(titulo_original)
            dados["titulo"] = titulo_corrigido
            match = REGEX_N_EDITAL.search(titulo_corrigido)
            if match:
                dados["n_edital"] = match.group(1).replace(" ", "")

        descricao_original = ""
        if metadados.get("descricao"):
            descricao_original = str(metadados["descricao"])
            dados["descricao"] = limpar_texto(corrigir_encoding(descricao_original), max_length=500)
        elif metadados.get("objeto"):
            descricao_original = str(metadados["objeto"])
            dados["descricao"] = limpar_texto(corrigir_encoding(descricao_original), max_length=500)

        # V11: Extrai data_publicacao e data_atualizacao
        if metadados.get("data_publicacao"):
            dados["data_publicacao"] = str(metadados["data_publicacao"]).split("T")[0]

        if metadados.get("data_atualizacao"):
            dados["data_atualizacao"] = str(metadados["data_atualizacao"]).split("T")[0]

        if metadados.get("data_inicio_propostas"):
            dados["data_leilao"] = str(metadados["data_inicio_propostas"]).split("T")[0]
        elif metadados.get("dataAberturaPropostas"):
            dados["data_leilao"] = str(metadados["dataAberturaPropostas"]).split("T")[0]
        elif descricao_original:
            data_extraida = extrair_data_de_texto(descricao_original)
            if data_extraida:
                dados["data_leilao"] = data_extraida

        urls_todas: List[str] = []
        texto_completo = ""

        if descricao_original:
            urls_todas.extend(extrair_urls_de_texto(descricao_original))
            texto_completo += descricao_original + " "

        objeto = str(metadados.get("objeto") or "")
        if objeto:
            urls_todas.extend(extrair_urls_de_texto(objeto))
            texto_completo += objeto + " "

        info = metadados.get("informacoes_complementares")
        if info is not None:
            info_str = str(info)
            urls_todas.extend(extrair_urls_de_texto(info_str))
            texto_completo += info_str + " "

        for file_info in metadados.get("files_meta", []) or []:
            titulo_file = str(file_info.get("titulo") or "")
            if titulo_file:
                urls_todas.extend(extrair_urls_de_texto(titulo_file))

        link_leiloeiro = encontrar_link_leiloeiro(urls_todas, texto_completo)
        if link_leiloeiro:
            dados["link_leiloeiro"] = link_leiloeiro

        link_pncp = metadados.get("link_pncp")
        if link_pncp:
            dados["link_pncp"] = str(link_pncp)

        orgao = metadados.get("orgao_nome")
        if orgao:
            dados["orgao"] = str(orgao)

        uf = metadados.get("uf")
        if uf:
            dados["uf"] = str(uf)

        municipio = metadados.get("municipio")
        if municipio:
            dados["cidade"] = str(municipio)

    except Exception as exc:
        print(f"[AVISO] Erro ao ler JSON: {exc}")

    return dados


def extrair_de_excel(pasta: Path) -> Dict[str, str]:
    dados: Dict[str, str] = {}
    arquivos = list(pasta.glob("*.xlsx")) + list(pasta.glob("*.xls")) + list(pasta.glob("*.csv"))
    if not arquivos:
        return dados

    for arquivo in arquivos:
        try:
            if arquivo.suffix.lower() == ".csv":
                df = pd.read_csv(arquivo, encoding="utf-8", nrows=50)
            else:
                df = pd.read_excel(arquivo, nrows=50)
            if df.empty:
                continue

            for col in df.columns:
                col_lower = str(col).lower()

                if "edital" in col_lower and "n_edital" not in dados:
                    valores = df[col].dropna().astype(str).tolist()
                    if valores:
                        dados["n_edital"] = valores[0]

                if any(kw in col_lower for kw in ["data", "leilao", "abertura"]) and "data_leilao" not in dados:
                    valores = df[col].dropna().astype(str).tolist()
                    if valores:
                        data = extrair_data_de_texto(valores[0])
                        if data:
                            dados["data_leilao"] = data

                if any(kw in col_lower for kw in ["desc", "objeto", "bem"]) and "descricao" not in dados:
                    valores = df[col].dropna().astype(str).tolist()
                    if valores:
                        dados["descricao"] = " | ".join(valores[:3])
        except Exception:
            continue

    return dados


def extrair_de_docx(pasta: Path) -> Dict[str, str]:
    dados: Dict[str, str] = {}
    if not DOCX_AVAILABLE:
        return dados

    arquivos = list(pasta.glob("*.docx")) + list(pasta.glob("*.doc"))
    for docx_path in arquivos:
        texto = extrair_texto_docx(docx_path)
        if not texto:
            continue

        if "n_edital" not in dados:
            match = REGEX_N_EDITAL.search(texto)
            if match:
                dados["n_edital"] = match.group(1).replace(" ", "")

        if "data_leilao" not in dados:
            data = extrair_data_de_texto(texto)
            if data:
                dados["data_leilao"] = data

        if "link_leiloeiro" not in dados:
            urls = extrair_urls_de_texto(texto)
            link = encontrar_link_leiloeiro(urls, texto)
            if link:
                dados["link_leiloeiro"] = link

    return dados


def extrair_de_path(pasta_edital: Path) -> Dict[str, str]:
    dados: Dict[str, str] = {}
    partes = pasta_edital.name.split("_")

    if len(partes) >= 3:
        dados["uf"] = partes[0].upper()
        dados["cidade"] = partes[1].replace("-", " ").title()
        dados["orgao"] = " ".join(partes[2:]).replace("-", " ").title()
    elif len(partes) >= 2:
        dados["uf"] = partes[0].upper()
        dados["cidade"] = partes[1].replace("-", " ").title()

    return dados


def extrair_de_pdf(pasta: Path) -> Dict[str, str]:
    dados: Dict[str, str] = {}
    pdfs = sorted(list(pasta.glob("*.pdf")))
    if not pdfs:
        return dados

    texto_completo = extrair_texto_pdfs(pdfs)
    if not texto_completo:
        return dados

    match = REGEX_N_EDITAL.search(texto_completo)
    if match:
        dados["n_edital"] = match.group(1).replace(" ", "")

    data_extraida = extrair_data_de_texto(texto_completo)
    if data_extraida:
        dados["data_leilao"] = data_extraida

    match_titulo = re.search(
        r"(?i)(?:objeto|finalidade).*?[:;]\s*(.{10,200}?)(?:\.|;|\n\n)",
        texto_completo,
        re.DOTALL,
    )
    if match_titulo:
        dados["titulo"] = match_titulo.group(1).strip()[:200]

    # V11: Extrai objeto_resumido
    objeto_resumido = extrair_objeto_resumido(texto_completo)
    if objeto_resumido != "N/D":
        dados["objeto_resumido"] = objeto_resumido

    urls = extrair_urls_de_texto(texto_completo)
    urls.extend(extrair_urls_contextuais(texto_completo))
    link = encontrar_link_leiloeiro(urls, texto_completo)
    if link:
        dados["link_leiloeiro"] = link

    return dados


def extrair_urls_de_todos_arquivos(pasta: Path) -> Tuple[List[str], str]:
    urls: List[str] = []
    texto: List[str] = []

    urls.extend(extrair_urls_de_docx(pasta))
    urls.extend(extrair_urls_de_xlsx(pasta))
    urls.extend(extrair_urls_de_zip(pasta))

    for docx_path in list(pasta.glob("*.docx")) + list(pasta.glob("*.doc")):
        doc_text = extrair_texto_docx(docx_path)
        if doc_text:
            texto.append(doc_text)

    pdfs = sorted(list(pasta.glob("*.pdf")))
    pdf_text = extrair_texto_pdfs(pdfs)
    if pdf_text:
        texto.append(pdf_text)

    texto_completo = " ".join(texto).strip()
    return list(dict.fromkeys(urls)), texto_completo


def pasta_parece_edital(pasta: Path) -> bool:
    if not pasta.is_dir():
        return False
    if (pasta / "metadados_pncp.json").exists():
        return True
    padroes = ["*.pdf", "*.docx", "*.doc", "*.xlsx", "*.xls", "*.xlsm", "*.csv", "*.zip"]
    return any(pasta.glob(p) for p in padroes)


def listar_pastas_editais(pasta_raiz: Path) -> List[Path]:
    if not pasta_raiz.exists():
        return []

    editais: List[Path] = []
    municipios = [d for d in pasta_raiz.iterdir() if d.is_dir()]

    for municipio in municipios:
        subdirs = [d for d in municipio.iterdir() if d.is_dir()]
        subdirs_editais = [d for d in subdirs if pasta_parece_edital(d)]

        if subdirs_editais:
            editais.extend(sorted(subdirs_editais))
            continue

        if pasta_parece_edital(municipio):
            editais.append(municipio)

    return editais


# ============================================================================
# V12: FUNÇÕES DE CORREÇÃO CRÍTICA
# ============================================================================

def detectar_leilao_presencial(modalidade: str, local_realizacao: str, descricao: str) -> bool:
    """Detecta se leilão é presencial baseado em múltiplas fontes"""
    texto = f"{modalidade} {local_realizacao} {descricao}".lower()
    return any(x in texto for x in ['presencial', 'sede', 'auditório', 'sala', 'comparecimento', 'local:'])


def validar_link_leiloeiro_v12(url: str, modalidade: str = "", local_realizacao: str = "", descricao: str = "") -> str:
    """
    V12 BUG #2: Valida link do leiloeiro com lógica condicional.
    - Leilão ONLINE: link obrigatório
    - Leilão PRESENCIAL: link pode ser "PRESENCIAL" (ausência válida)
    - Rejeita domínios de email (hotmail, yahoo, gmail, etc)
    """
    eh_presencial = detectar_leilao_presencial(modalidade, local_realizacao, descricao)

    if not url or url.strip() == '' or url == 'N/D':
        return "PRESENCIAL" if eh_presencial else "N/D"

    # Normalizar URL
    url_norm = url.strip().lower()
    if not url_norm.startswith(('http://', 'https://')):
        url_norm = 'https://' + url_norm

    # Extrair domínio
    try:
        from urllib.parse import urlparse
        dominio = urlparse(url_norm).netloc.replace('www.', '')
    except:
        return "PRESENCIAL" if eh_presencial else "N/D"

    # Verificar se domínio é inválido (email)
    if dominio in DOMINIOS_INVALIDOS:
        return "PRESENCIAL" if eh_presencial else "N/D"

    # Verificar se parece um site real
    if '.' not in dominio or len(dominio) < 4:
        return "PRESENCIAL" if eh_presencial else "N/D"

    return url_norm


def extrair_componentes_pncp_v12(json_data: dict, path_metadata: dict, link_pncp_atual: str = "") -> tuple:
    """
    V12 BUG #3: Extrai CNPJ, ANO e SEQUENCIAL das fontes disponíveis.
    Incluindo extração do link_pncp atual se estiver no formato antigo.
    """
    # Tentar extrair do link_pncp atual (formato antigo: /CNPJ-MODALIDADE-SEQUENCIAL/ANO)
    if link_pncp_atual and 'pncp.gov.br' in link_pncp_atual:
        # Formato: https://pncp.gov.br/app/editais/04302189000128-1-000019/2025
        match = re.search(r'/editais/([\d]+)[-\d]+-([\d]+)/(\d{4})', link_pncp_atual)
        if match:
            cnpj_link = match.group(1)
            sequencial_link = match.group(2)
            ano_link = match.group(3)
            return cnpj_link, ano_link, sequencial_link

    # CNPJ do órgão
    cnpj = (
        json_data.get('orgaoEntidade', {}).get('cnpj') if isinstance(json_data.get('orgaoEntidade'), dict) else None or
        json_data.get('cnpjOrgao') or
        json_data.get('cnpj') or
        json_data.get('orgao_cnpj') or
        path_metadata.get('cnpj') or
        ''
    )

    # ANO da publicação
    ano = (
        json_data.get('anoCompra') or
        str(json_data.get('dataPublicacaoPncp', ''))[:4] or
        str(json_data.get('data_publicacao', ''))[:4] or
        path_metadata.get('ano') or
        ''
    )

    # SEQUENCIAL da compra
    sequencial = (
        json_data.get('sequencialCompra') or
        json_data.get('numeroCompra') or
        json_data.get('sequencial') or
        path_metadata.get('sequencial') or
        ''
    )

    return cnpj, ano, sequencial



def montar_link_pncp_v12(cnpj: str, ano: str, sequencial: str) -> str:
    """
    V12 BUG #3 CRÍTICO: Monta link PNCP no formato OFICIAL CORRETO.

    Formato: https://pncp.gov.br/app/editais/{CNPJ}/{ANO}/{SEQUENCIAL}
    Exemplo: https://pncp.gov.br/app/editais/88150495000186/2025/000490
    """
    # Limpar CNPJ (apenas números, 14 dígitos)
    cnpj_limpo = re.sub(r'\D', '', str(cnpj))
    if len(cnpj_limpo) != 14:
        return "N/D"

    # Limpar e validar ano (4 dígitos)
    ano_limpo = re.sub(r'\D', '', str(ano))
    if len(ano_limpo) != 4:
        return "N/D"

    # Limpar sequencial (remover zeros à esquerda mas manter se for só zeros)
    sequencial_limpo = re.sub(r'\D', '', str(sequencial))
    if not sequencial_limpo:
        return "N/D"

    # Remover zeros à esquerda
    sequencial_limpo = sequencial_limpo.lstrip('0') or '0'

    # FORMATO CORRETO: /CNPJ/ANO/SEQUENCIAL (sem zeros à esquerda no sequencial)
    return f"https://pncp.gov.br/app/editais/{cnpj_limpo}/{ano_limpo}/{sequencial_limpo}"



def extrair_tags_inteligente_v12(descricao: str, pdf_text: str, titulo: str) -> str:
    """
    V12 BUG #4: Extrai tags específicas do conteúdo real do edital.
    Retorna lista de tags separadas por vírgula.
    """
    texto_completo = f"{titulo} {descricao} {pdf_text[:3000]}".lower()
    tags_encontradas = set()

    for tag, palavras_chave in MAPA_TAGS.items():
        for palavra in palavras_chave:
            if palavra.lower() in texto_completo:
                tags_encontradas.add(tag)
                break

    # Garantir pelo menos uma tag
    if not tags_encontradas:
        if 'veículo' in texto_completo or 'veiculo' in texto_completo:
            tags_encontradas.add('veiculo')
        if 'leilão' in texto_completo or 'leilao' in texto_completo:
            tags_encontradas.add('leilao')

    if not tags_encontradas:
        return "sem_classificacao"

    return ','.join(sorted(tags_encontradas))


def extrair_titulo_inteligente_v12(pdf_text: str, json_data: dict, n_edital: str) -> str:
    """
    V12 BUG #5: Extrai título da primeira linha do PDF, limitado a 100 caracteres.
    """
    # FONTE 1: Primeira linha do PDF
    if pdf_text:
        linhas = pdf_text.strip().split('\n')
        for linha in linhas[:10]:
            linha_limpa = linha.strip()
            if len(linha_limpa) > 20 and not linha_limpa.replace(' ', '').isdigit():
                ignorar = ['ministério', 'secretaria', 'governo', 'estado', 'página', 'pag.', 'poder executivo']
                if not any(ig in linha_limpa.lower() for ig in ignorar):
                    return linha_limpa[:100]

    # FONTE 2: Objeto do JSON PNCP
    objeto = json_data.get('objetoCompra', '') or json_data.get('objeto', '')
    if objeto and len(objeto) > 20:
        return objeto[:100]

    # FONTE 3: Fallback
    return f"Edital nº {n_edital}" if n_edital else "Edital sem identificação"


def extrair_modalidade_v12(json_data: dict, pdf_text: str, descricao: str = "") -> str:
    """V12: Retorna: ONLINE | PRESENCIAL | HÍBRIDO | N/D"""
    texto = f"{json_data.get('modalidadeNome', '')} {pdf_text[:2000]} {descricao}".lower()

    tem_online = any(x in texto for x in ['eletrônico', 'eletronico', 'online', 'internet', 'virtual', 'plataforma digital'])
    tem_presencial = any(x in texto for x in ['presencial', 'sede', 'auditório', 'sala', 'comparecimento', 'local:'])

    if tem_online and tem_presencial:
        return "HÍBRIDO"
    elif tem_online:
        return "ONLINE"
    elif tem_presencial:
        return "PRESENCIAL"
    return "N/D"


def formatar_valor_br(valor_raw) -> str:
    """Formata valor numérico para formato brasileiro R$ X.XXX,XX"""
    if not valor_raw:
        return "N/D"

    try:
        valor_float = float(valor_raw)
        return f"R$ {valor_float:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.')
    except:
        return "N/D"


def extrair_valor_estimado_v12(json_data: dict, pdf_text: str, arquivo_origem: str = "") -> str:
    """
    V12 + API: Extrai valor estimado/mínimo do leilão.

    Ordem de prioridade:
    0. API PNCP COMPLETA (valorTotalEstimado) ← NOVO! PRIORIDADE MÁXIMA!
    1. JSON PNCP local
    2. PDF com regex
    """

    # FONTE 0: API PNCP COMPLETA - PRIORIDADE ABSOLUTA!
    if arquivo_origem:
        cnpj, ano, sequencial = extrair_componentes_do_path_edital(arquivo_origem)
        if cnpj and ano and sequencial:
            api_data = buscar_api_pncp_completa(cnpj, ano, sequencial)
            if api_data:
                valor_api = api_data.get('valorTotalEstimado')
                if valor_api:
                    valor_formatado = formatar_valor_br(valor_api)
                    if valor_formatado != "N/D":
                        return valor_formatado

    # FONTE 1: JSON PNCP local (fallback)
    valor = json_data.get('valorTotalEstimado') or json_data.get('valorEstimado') or json_data.get('valorTotal')
    if valor:
        valor_formatado = formatar_valor_br(valor)
        if valor_formatado != "N/D":
            return valor_formatado

    # FONTE 2: PDF com regex (último recurso)
    padrao = r'(?:valor|lance|mínimo|avaliação|avaliacao|estimado)[:\s]*R?\$?\s*([\d.,]+)'
    match = re.search(padrao, pdf_text[:3000], re.IGNORECASE)
    if match:
        valor_str = match.group(1).replace('.', '').replace(',', '.')
        valor_formatado = formatar_valor_br(valor_str)
        if valor_formatado != "N/D":
            return valor_formatado

    return "N/D"


def extrair_quantidade_itens_v12(json_data: dict, pdf_text: str) -> str:
    """V12: Extrai quantidade de itens/lotes do leilão."""
    qtd = json_data.get('quantidadeItens') or json_data.get('numeroItens')
    if qtd:
        return str(qtd)

    lotes = len(re.findall(r'\bLOTE\s*\d+', pdf_text[:5000], re.IGNORECASE))
    if lotes > 0:
        return str(lotes)

    itens = len(re.findall(r'\bITEM\s*\d+', pdf_text[:5000], re.IGNORECASE))
    if itens > 0:
        return str(itens)

    return "N/D"


def extrair_nome_leiloeiro_v12(json_data: dict, pdf_text: str) -> str:
    """V12: Extrai nome do leiloeiro oficial."""
    leiloeiro = json_data.get('nomeLeiloeiro') or json_data.get('leiloeiro') or json_data.get('responsavel')
    if leiloeiro:
        return str(leiloeiro).strip()[:100]

    padrao = r'(?:leiloeiro|leiloeira)[:\s]*(?:oficial|público|a)?\s*[:\s]*([A-ZÀ-Ú][a-zà-ú]+(?:\s+[A-ZÀ-Ú][a-zà-ú]+){1,4})'
    match = re.search(padrao, pdf_text[:3000])
    if match:
        return match.group(1).strip()[:100]

    padrao2 = r'(?:responsável|responsavel)[:\s]*(?:pelo\s+leilão|pela\s+venda)?\s*[:\s]*([A-ZÀ-Ú][a-zà-ú]+(?:\s+[A-ZÀ-Ú][a-zà-ú]+){1,4})'
    match2 = re.search(padrao2, pdf_text[:3000], re.IGNORECASE)
    if match2:
        return match2.group(1).strip()[:100]

    return "N/D"


def extrair_componentes_do_path_edital(arquivo_origem: str) -> tuple:
    """
    Extrai CNPJ, ANO, SEQUENCIAL do caminho do edital.
    Exemplo: "AM_MANAUS/2025-11-21_S60_04312641000132-1-000097-2025"
    Retorna: (cnpj, ano, sequencial) ou (None, None, None)
    """
    # Padrão: _{CNPJ}-{CODIGO}-{SEQUENCIAL}-{ANO}
    match = re.search(r'_(\d{14})-\d+-(\d+)-(\d{4})$', arquivo_origem)
    if match:
        cnpj = match.group(1)
        sequencial = match.group(2).lstrip('0') or '0'
        ano = match.group(3)
        return (cnpj, ano, sequencial)
    return (None, None, None)


def buscar_api_pncp_completa(cnpj: str, ano: str, sequencial: str) -> dict:
    """
    Busca o JSON COMPLETO da API do PNCP.
    Endpoint: https://pncp.gov.br/api/consulta/v1/orgaos/{cnpj}/compras/{ano}/{sequencial}

    Retorna dict com dados da API ou dict vazio se houver erro.
    """
    if not cnpj or not ano or not sequencial:
        return {}

    url = f"https://pncp.gov.br/api/consulta/v1/orgaos/{cnpj}/compras/{ano}/{sequencial}"

    try:
        import requests
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=10)
        response.raise_for_status()
        return response.json()
    except Exception:
        # Silenciosamente retorna vazio se API falhar
        return {}


def extrair_data_leilao_cascata_v12(json_data: dict, pdf_text: str, excel_data: dict, descricao: str = "", arquivo_origem: str = "") -> str:
    """
    V12 BUG #1 CRÍTICO: Cascata de extração para data_leilao.
    SEM ESSA DATA NÃO EXISTE ACHE SUCATAS!

    Ordem de prioridade:
    0. API PNCP COMPLETA (dataAberturaProposta) ← NOVO! PRIORIDADE MÁXIMA!
    1. JSON PNCP local (campos de data)
    2. DESCRIÇÃO (primeiro, pois vem do JSON)
    3. Excel/CSV anexo
    4. PDF com padrões agressivos
    """
    from local_auditor_v12_final import formatar_data_br

    # FONTE 0: API PNCP COMPLETA - PRIORIDADE ABSOLUTA!
    if arquivo_origem:
        cnpj, ano, sequencial = extrair_componentes_do_path_edital(arquivo_origem)
        if cnpj and ano and sequencial:
            api_data = buscar_api_pncp_completa(cnpj, ano, sequencial)
            if api_data:
                # Campo dataAberturaProposta é a data do leilão
                data_abertura = api_data.get('dataAberturaProposta')
                if data_abertura:
                    data_formatada = formatar_data_br(data_abertura)
                    if data_formatada != "N/D":
                        return data_formatada

    # FONTE 1: JSON PNCP local - Prioridade máxima
    campos_json = [
        'dataAberturaProposta', 'dataAberturaPropostas', 'data_inicio_propostas',
        'dataEncerramentoProposta', 'dataInicioVigencia',
        'dataFimProposta', 'dataFimPropostas',
        'dataRealizacao', 'dataLeilao'
    ]
    for campo in campos_json:
        if json_data.get(campo):
            data_formatada = formatar_data_br(json_data[campo])
            if data_formatada != "N/D":
                return data_formatada

    # FONTE 2: DESCRIÇÃO (texto do JSON) - MUITO IMPORTANTE!
    if descricao:
        # Padrões específicos de leilão
        padroes_desc = [
            r'(?:leil[ãa]o.*?dia|dia.*?leil[ãa]o).*?(\d{1,2}[/.-]\d{1,2}[/.-]20\d{2})',
            r'(?:realizar[aá]|ocorrer[aá]).*?(\d{1,2}[/.-]\d{1,2}[/.-]20\d{2})',
            r'(?:data|dia).*?(\d{1,2}[/.-]\d{1,2}[/.-]20\d{2}).*?(?:leil[ãa]o|hasta|arremata)',
            r'(\d{1,2}[/.-]\d{1,2}[/.-]20\d{2}).*?(?:às|as|hora|h)\s*\d{1,2}',
        ]

        for padrao in padroes_desc:
            match = re.search(padrao, descricao[:2000], re.IGNORECASE | re.DOTALL)
            if match:
                data_formatada = formatar_data_br(match.group(1))
                if data_formatada != "N/D":
                    return data_formatada

    # FONTE 3: Excel/CSV anexo
    if excel_data:
        colunas_data = [c for c in excel_data.keys() if 'data' in c.lower() or 'leilao' in c.lower()]
        for col in colunas_data:
            if excel_data[col] and excel_data[col] != 'N/D':
                data_formatada = formatar_data_br(excel_data[col])
                if data_formatada != "N/D":
                    return data_formatada

    # FONTE 4: PDF - PADRÕES MUITO AGRESSIVOS
    if pdf_text:
        # Lista de padrões ordenados por especificidade
        padroes_pdf = [
            # Padrão 1: "data do leilão: DD/MM/YYYY"
            r'(?:data\s*(?:do|de)\s*leil[ãa]o)[:\s]*(\d{1,2}[/.-]\d{1,2}[/.-]20\d{2})',

            # Padrão 2: "leilão dia DD/MM/YYYY"
            r'(?:leil[ãa]o|hasta|pregão).*?(?:dia|data)[:\s]*(\d{1,2}[/.-]\d{1,2}[/.-]20\d{2})',

            # Padrão 3: "será realizado em DD/MM/YYYY"
            r'(?:será|ser[aá])\s*realizado.*?(\d{1,2}[/.-]\d{1,2}[/.-]20\d{2})',

            # Padrão 4: "ocorrerá no dia DD/MM/YYYY"
            r'(?:ocorrer[aá]|realizar[aá]).*?(?:dia|em)[:\s]*(\d{1,2}[/.-]\d{1,2}[/.-]20\d{2})',

            # Padrão 5: "abertura" ou "sessão" com data
            r'(?:abertura|sessão|sess[ãa]o)[:\s]*.*?(\d{1,2}[/.-]\d{1,2}[/.-]20\d{2})',

            # Padrão 6: Data seguida de horário (forte indicativo)
            r'(\d{1,2}[/.-]\d{1,2}[/.-]20\d{2})[,\s]*(?:às|as|[àa]s|hora)[:\s]*\d{1,2}',

            # Padrão 7: "dia DD/MM/YYYY às HH"
            r'dia\s*(\d{1,2}[/.-]\d{1,2}[/.-]20\d{2})\s*[àa]s?\s*\d{1,2}',

            # Padrão 8: Formato de convocação
            r'(?:comparecer|participar).*?(\d{1,2}[/.-]\d{1,2}[/.-]20\d{2})',

            # Padrão 9: Data em contexto de evento
            r'(?:no\s*dia|na\s*data)[:\s]*(\d{1,2}[/.-]\d{1,2}[/.-]20\d{2})',

            # Padrão 10: AGRESSIVO - Primeira data válida após palavra-chave
            r'(?:leil[ãa]o|hasta|arremata|aliena|pregão|venda)(?:.*?(\d{1,2}[/.-]\d{1,2}[/.-]20\d{2})){1}',

            # Padrão 11: MUITO AGRESSIVO - Qualquer data DD/MM/20YY nos primeiros 3000 chars
            r'(\d{1,2}[/.-]\d{1,2}[/.-]20\d{2})',
        ]

        # Buscar nos primeiros 5000 caracteres (geralmente a data está no início)
        texto_busca = pdf_text[:5000]

        for padrao in padroes_pdf:
            matches = re.finditer(padrao, texto_busca, re.IGNORECASE | re.DOTALL)
            for match in matches:
                data_str = match.group(1)
                data_formatada = formatar_data_br(data_str)
                if data_formatada != "N/D":
                    # Validar se é uma data futura ou recente (não muito antiga)
                    try:
                        from datetime import datetime
                        partes = data_formatada.split('/')
                        data_obj = datetime(int(partes[2]), int(partes[1]), int(partes[0]))
                        # Aceitar datas de 2020 em diante
                        if data_obj.year >= 2020:
                            return data_formatada
                    except:
                        # Se falhar a validação, ainda retorna
                        return data_formatada

    return "N/D"



def extrair_data_atualizacao_cascata_v12(json_data: dict) -> str:
    """V12 BUG #1: Data de atualização vem EXCLUSIVAMENTE do JSON PNCP."""
    campos = ['dataAtualizacao', 'dataModificacao', 'dataUltimaAtualizacao', 'data_atualizacao']
    for campo in campos:
        if json_data.get(campo):
            data_formatada = formatar_data_br(json_data[campo])
            if data_formatada != "N/D":
                return data_formatada

    # Fallback: usar data de publicação
    if json_data.get('dataPublicacaoPncp') or json_data.get('data_publicacao'):
        data_pub = json_data.get('dataPublicacaoPncp') or json_data.get('data_publicacao')
        data_formatada = formatar_data_br(data_pub)
        if data_formatada != "N/D":
            return data_formatada

    return "N/D"


def formatar_data_br(data_raw) -> str:
    """Formata qualquer formato de data para DD/MM/YYYY brasileiro"""
    if not data_raw or data_raw == "N/D":
        return "N/D"

    data_str = str(data_raw).strip()

    # Já está no formato DD/MM/YYYY
    if re.match(r'^\d{2}/\d{2}/\d{4}$', data_str):
        return data_str

    # Formato ISO (YYYY-MM-DD ou YYYY-MM-DDTHH:MM:SS)
    match = re.match(r'^(\d{4})-(\d{2})-(\d{2})', data_str)
    if match:
        ano, mes, dia = match.groups()
        return f"{dia}/{mes}/{ano}"

    # Formato DD-MM-YYYY ou DD.MM.YYYY
    match = re.match(r'^(\d{2})[-.](\d{2})[-.](\d{4})$', data_str)
    if match:
        dia, mes, ano = match.groups()
        return f"{dia}/{mes}/{ano}"

    # Timestamp pandas ou datetime
    try:
        from datetime import datetime
        if isinstance(data_raw, (pd.Timestamp, datetime)):
            return data_raw.strftime('%d/%m/%Y')
    except:
        pass

    return "N/D"


# ============================================================================
# FIM FUNÇÕES V12
# ============================================================================


def processar_edital(pasta_edital: Path) -> Optional[Dict[str, str]]:
    print(f"[INFO] Processando: {pasta_edital.relative_to(PASTA_EDITAIS)}")

    if not pasta_edital.exists() or not pasta_edital.is_dir():
        print("  [ERRO] Pasta do edital invalida")
        return None

    resultado = ResultadoEdital(arquivo_origem=str(pasta_edital.relative_to(PASTA_EDITAIS)))
    dados_finais = resultado.as_dict()

    # Extração básica de todas as fontes
    fontes: List[Tuple[str, Dict[str, str]]] = [
        ("JSON", extrair_de_metadados_json(pasta_edital)),
        ("Excel", extrair_de_excel(pasta_edital)),
        ("DOCX", extrair_de_docx(pasta_edital)),
        ("Path", extrair_de_path(pasta_edital)),
        ("PDF", extrair_de_pdf(pasta_edital)),
    ]

    # Guardar JSON, Excel e Path para V12
    dados_json = fontes[0][1]
    json_data = dados_json.get('_json_raw', {}) if isinstance(dados_json.get('_json_raw'), dict) else {}
    excel_data = fontes[1][1]
    path_data = fontes[3][1]

    campos_encontrados: List[str] = []
    for nome_fonte, dados_fonte in fontes:
        for campo, valor in dados_fonte.items():
            if campo not in dados_finais or campo == '_json_raw':
                continue
            if dados_finais[campo] == "N/D" and valor and valor != "N/D":
                dados_finais[campo] = valor
                campos_encontrados.append(f"{campo}({nome_fonte})")

    # Busca profunda de link_leiloeiro se necessário
    if dados_finais["link_leiloeiro"] == "N/D":
        urls_todos, texto_todos = extrair_urls_de_todos_arquivos(pasta_edital)
        link_encontrado = encontrar_link_leiloeiro(urls_todos, texto_todos)
        if link_encontrado:
            dados_finais["link_leiloeiro"] = link_encontrado
            campos_encontrados.append("link_leiloeiro(V10_DEEP)")

    # ========================================================================
    # V12: APLICAR TODAS AS CORREÇÕES CRÍTICAS
    # ========================================================================

    # V12: Extrair texto PDF para análise detalhada
    pdfs = sorted(pasta_edital.glob("*.pdf"))
    pdf_text = extrair_texto_pdfs(pdfs) if pdfs else ""

    # BUG #1: Datas com cascata inteligente
    dados_finais["data_leilao"] = extrair_data_leilao_cascata_v12(
        json_data, pdf_text, excel_data, dados_finais.get("descricao", ""), dados_finais.get("arquivo_origem", "")
    )
    dados_finais["data_atualizacao"] = extrair_data_atualizacao_cascata_v12(json_data)

    # BUG #2: Validar link_leiloeiro (rejeita emails, aceita PRESENCIAL)
    dados_finais["link_leiloeiro"] = validar_link_leiloeiro_v12(
        dados_finais["link_leiloeiro"],
        dados_finais.get("modalidade_leilao", ""),
        "",
        dados_finais.get("descricao", "")
    )

    # BUG #3: Corrigir link_pncp para formato /CNPJ/ANO/SEQUENCIAL
    # CRÍTICO: SEMPRE sobrescrever o link_pncp, mesmo que já exista!
    cnpj, ano, sequencial = extrair_componentes_pncp_v12(
        json_data, path_data, dados_finais.get("link_pncp", "")
    )
    if cnpj and ano and sequencial:
        link_pncp_correto = montar_link_pncp_v12(cnpj, ano, sequencial)
        if link_pncp_correto != "N/D":
            dados_finais["link_pncp"] = link_pncp_correto  # SEMPRE sobrescrever!
            campos_encontrados.append("link_pncp(V12_CORRIGIDO)")

    # BUG #4: Tags inteligentes baseadas no conteúdo
    dados_finais["tags"] = extrair_tags_inteligente_v12(
        dados_finais.get("descricao", ""),
        pdf_text,
        dados_finais.get("titulo", "")
    )

    # BUG #5: Título da primeira linha do PDF
    dados_finais["titulo"] = extrair_titulo_inteligente_v12(
        pdf_text,
        json_data,
        dados_finais.get("n_edital", "")
    )

    # NOVOS CAMPOS V12
    dados_finais["modalidade_leilao"] = extrair_modalidade_v12(
        json_data, pdf_text, dados_finais.get("descricao", "")
    )
    dados_finais["valor_estimado"] = extrair_valor_estimado_v12(json_data, pdf_text, dados_finais.get("arquivo_origem", ""))
    dados_finais["quantidade_itens"] = extrair_quantidade_itens_v12(json_data, pdf_text)
    dados_finais["nome_leiloeiro"] = extrair_nome_leiloeiro_v12(json_data, pdf_text)

    # V11: Gera id_interno baseado no link_pncp (após correção V12)
    dados_finais["id_interno"] = gerar_id_interno(dados_finais["link_pncp"])

    # ========================================================================
    # FIM V12
    # ========================================================================

    if campos_encontrados:
        preview = campos_encontrados[:5]
        suffix = "..." if len(campos_encontrados) > 5 else ""
        print(f"  [OK] Extraido: {', '.join(preview)}{suffix}")
    else:
        print("  [AVISO] Nenhum dado extraido")

    return dados_finais


def varrer_editais() -> List[Dict[str, str]]:
    if not PASTA_EDITAIS.exists():
        print(f"[ERRO] Diretorio {PASTA_EDITAIS} nao encontrado")
        return []

    print(f"[INFO] Varrendo editais em: {PASTA_EDITAIS}")

    pastas_editais = listar_pastas_editais(PASTA_EDITAIS)
    print(f"[INFO] Total de editais encontrados: {len(pastas_editais)}\n")

    resultados: List[Dict[str, str]] = []
    for i, pasta in enumerate(pastas_editais, 1):
        try:
            print(f"[{i}/{len(pastas_editais)}] ", end="")
            dados = processar_edital(pasta)
            if dados:
                resultados.append(dados)
        except Exception as exc:
            print(f"[ERRO] Falha ao processar {pasta}: {exc}")

    return resultados


def extrair_n_pncp(link_pncp: str) -> str:
    """Extrai n_pncp do link_pncp, retorna N/D se não encontrar"""
    if not link_pncp or link_pncp == "N/D":
        return "N/D"

    # Tenta extrair do padrão: /editais/{cnpj}/{ano}/{sequencial}
    match = re.search(r"/editais/([^/]+/\d{4}/\d+)", link_pncp)
    if match:
        return match.group(1)

    # Fallback: tenta extrair qualquer padrão numérico após /editais/
    match = re.search(r"/editais/([^/\s]+)", link_pncp)
    if match:
        return match.group(1)

    return "N/D"


def salvar_resultado_final(dados: List[Dict[str, str]]) -> None:
    """Gera arquivo RESULTADO_FINAL.xlsx com todos os campos na ordem especificada"""
    if not dados:
        print("[ERRO] Nenhum dado para salvar")
        return

    # Adiciona campo n_pncp extraído do link_pncp
    for item in dados:
        item["n_pncp"] = extrair_n_pncp(item.get("link_pncp", "N/D"))

    df = pd.DataFrame(dados)

    # Reordena colunas conforme especificado
    colunas_ordenadas = [
        "id_interno", "orgao", "uf", "cidade", "n_pncp", "n_edital",
        "data_publicacao", "data_atualizacao", "data_leilao",
        "titulo", "descricao", "tags", "link_pncp", "link_leiloeiro",
        "objeto_resumido", "arquivo_origem"
    ]

    # Adiciona colunas faltantes com N/D
    for col in colunas_ordenadas:
        if col not in df.columns:
            df[col] = "N/D"

    df = df[colunas_ordenadas]

    output_path = PASTA_RAIZ / "RESULTADO_FINAL.xlsx"
    df.to_excel(output_path, index=False, engine="openpyxl")

    print(f"\n{'='*60}")
    print(f"[OK] RESULTADO FINAL GERADO: {output_path}")
    print(f"[OK] TOTAL DE REGISTROS: {len(dados)}")
    print(f"{'='*60}\n")

    # Exibe estatísticas de preenchimento
    print("=" * 60)
    print("VALIDAÇÃO RESULTADO_FINAL.xlsx")
    print("=" * 60)
    print(f"{'[CAMPO]':<25} {'[PREENCHIDOS]':<20} {'[TAXA]':<10}")
    print("-" * 60)

    for col in colunas_ordenadas:
        if col == "arquivo_origem":
            continue
        total = len(df)
        nao_vazio = df[col].notna() & (df[col] != "") & (df[col] != "N/D")
        preenchidos = int(nao_vazio.sum())
        taxa = (preenchidos / total * 100) if total > 0 else 0

        status = "[OK]" if taxa >= 80 else "[PARCIAL]" if taxa >= 50 else "[BAIXO]"
        print(f"{col:<25} {preenchidos:3d}/{total:3d} ({taxa:5.1f}%)    {status}")

    print("=" * 60 + "\n")


def salvar_csv(dados: List[Dict[str, str]]) -> None:
    if not dados:
        print("[AVISO] Nenhum dado para salvar")
        return

    df = pd.DataFrame(dados)
    df.to_csv(CSV_OUTPUT, index=False, encoding="utf-8-sig")

    print(f"\n[OK] Arquivo gerado: {CSV_OUTPUT}")
    print(f"[OK] Total de editais processados: {len(dados)}")

    print("\n" + "=" * 60)
    print("ESTATISTICAS DE PREENCHIMENTO - V10")
    print("=" * 60)

    for col in df.columns:
        if col == "arquivo_origem":
            continue
        total = len(df)
        nao_vazio = df[col].notna() & (df[col] != "") & (df[col] != "N/D")
        preenchidos = int(nao_vazio.sum())
        taxa = (preenchidos / total * 100) if total > 0 else 0

        status = "[OK]" if taxa == 100 else "[PARCIAL]" if taxa >= 50 else "[BAIXO]"
        barra = "#" * int(taxa / 5)
        print(f"{status} {col:20s} {preenchidos:3d}/{total:3d} ({taxa:5.1f}%) {barra}")


def main() -> None:
    print("\n" + "=" * 60)
    print("ACHE SUCATAS DaaS - AUDITOR V13 - INTEGRAÇÃO SUPABASE")
    print("=" * 60)
    print("\nV13 - NOVIDADES:")
    print("[+] Integração com Supabase PostgreSQL")
    print("[+] RLS (Row Level Security) ativado")
    print("[+] Dual storage: Supabase + CSV/XLSX backup")
    print("[+] API PNCP FONTE 0 (100% data_leilao + valor_estimado)")
    print("\nV12 - Mantido:")
    print("[+] Suporte a .leilao.br")
    print("[+] Datas extensas (14 de janeiro de 2026)")
    print("[+] Correções críticas (links, tags, títulos)")
    print("\nCascata: JSON -> Excel -> DOCX -> Path -> PDF")
    print("=" * 60 + "\n")

    if not DOCX_AVAILABLE:
        print("[AVISO] python-docx nao instalado. DOCX sera ignorado.\n")

    # V13: Inicializar Supabase
    print("[V13] Inicializando Supabase Repository...")
    supabase_repo = SupabaseRepository(enable_supabase=True)

    if supabase_repo.enable_supabase:
        print("[OK] Supabase conectado")
        count = supabase_repo.contar_editais()
        print(f"[INFO] Editais já no banco: {count}\n")
    else:
        print("[AVISO] Supabase desabilitado - apenas CSV/XLSX local\n")

    # Processar editais
    resultados = varrer_editais()

    # V13: Persistir no Supabase
    if supabase_repo.enable_supabase and resultados:
        print(f"\n{'='*60}")
        print(f"[V13] PERSISTINDO NO SUPABASE")
        print(f"{'='*60}")
        print(f"Total de editais: {len(resultados)}")

        sucessos = 0
        falhas = 0

        for i, edital in enumerate(resultados, 1):
            sucesso = supabase_repo.inserir_edital(edital)
            if sucesso:
                sucessos += 1
            else:
                falhas += 1

            # Progress
            if i % 20 == 0 or i == len(resultados):
                print(f"  [{i}/{len(resultados)}] OK: {sucessos}, Falhas: {falhas}")

        print(f"\n[OK] Supabase: {sucessos} inseridos/atualizados, {falhas} falhas")

        # Contar total final
        count_final = supabase_repo.contar_editais()
        print(f"[INFO] Total no banco após processamento: {count_final}")
        print(f"{'='*60}\n")

    # Salvar backups locais (sempre)
    salvar_csv(resultados)
    salvar_resultado_final(resultados)

    print("\n" + "=" * 60)
    print("PROCESSAMENTO CONCLUIDO - V13")
    print("=" * 60 + "\n")


if __name__ == "__main__":
    main()
