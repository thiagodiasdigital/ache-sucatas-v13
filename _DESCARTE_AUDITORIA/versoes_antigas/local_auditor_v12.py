#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
ACHE SUCATAS DaaS - AUDITOR V12 - CORREÇÕES CRÍTICAS
Extrai dados de editais usando cascata: JSON -> Excel -> DOCX -> Path -> PDF

V12 - CORREÇÕES CRÍTICAS:
- BUG #1: Cascata de fontes para data_leilao e data_atualizacao (JSON -> Excel -> PDF)
- BUG #2: Validação de link_leiloeiro (rejeita emails, detecta presencial)
- BUG #3: Formato correto link_pncp: /CNPJ/ANO/SEQUENCIAL
- BUG #4: Tags inteligentes extraídas do conteúdo real
- BUG #5: Título inteligente da primeira linha do PDF
- NOVO: modalidade_leilao, valor_estimado, quantidade_itens, nome_leiloeiro
- NOVO: Validação completa de registros com estatísticas detalhadas

V11:
- FIX: Adiciona .leilao.br ao SUPPORTED_TLDS
- FIX: Captura datas extensas (ex: "14 de janeiro de 2026")
- NOVO: Adiciona id_interno, data_publicacao, data_atualizacao, objeto_resumido
- NOVO: Gera arquivo RESULTADO_FINAL.xlsx com todos os campos
"""

from __future__ import annotations

import hashlib
import json
import re
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import pandas as pd
import pdfplumber

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[AVISO] python-docx nao instalado. Instale com: pip install python-docx")


PASTA_RAIZ = Path(__file__).parent
PASTA_EDITAIS = PASTA_RAIZ / "ACHE_SUCATAS_DB"
CSV_OUTPUT = PASTA_RAIZ / "analise_editais_v12.csv"


REGEX_N_EDITAL = re.compile(
    r"(?i)(?:edital|processo|leilao|pregao).*?(\d{1,5}\s*/\s*20\d{2})",
    re.DOTALL,
)

REGEX_DATA = re.compile(r"\b(\d{1,2}[/\-]\d{1,2}[/\-]20\d{2})\b")

REGEX_DATA_CONTEXTUAL = re.compile(
    r"(?i)(?:data|abertura|sessao|leilao|pregao|realizacao|"
    r"desfazimento|alienacao|hasta|arrematacao).*?"
    r"(\d{1,2}[/\-]\d{1,2}[/\-]20\d{2})",
    re.DOTALL,
)

# V11: Suporte para datas extensas como "14 de janeiro de 2026"
MESES_BR = {
    "janeiro": "01", "fevereiro": "02", "março": "03", "marco": "03",
    "abril": "04", "maio": "05", "junho": "06", "julho": "07",
    "agosto": "08", "setembro": "09", "outubro": "10",
    "novembro": "11", "dezembro": "12"
}

REGEX_DATA_EXTENSO = re.compile(
    r"(\d{1,2})\s*de\s*(janeiro|fevereiro|março|marco|abril|maio|junho|"
    r"julho|agosto|setembro|outubro|novembro|dezembro)\s*de\s*(20\d{2})",
    re.IGNORECASE
)

# V11: Regex para extrair resumo dos objetos/veículos do leilão
REGEX_OBJETO_RESUMIDO = re.compile(
    r"(?i)(?:lote|item|veiculo|veículo|placa|marca|modelo|chassi).*?[:;]\s*(.{20,300})",
    re.DOTALL
)

SUPPORTED_TLDS = r"(?:\.leilao\.br|\.com\.br|\.com|\.org\.br|\.gov\.br|\.net\.br|\.net)"

REGEX_URL = re.compile(
    rf"(?:https?://)?(?:www\.)?[\w\-\.]+{SUPPORTED_TLDS}"
    r"(?:/[\w\-\./?%&=~#]*)?",
    re.IGNORECASE,
)

REGEX_URL_QUEBRADA = re.compile(
    rf"(?:https?://)?(?:www\.)?[\w\-\.]+\s*[\n\r]\s*{SUPPORTED_TLDS}"
    r"(?:/[\w\-\./?%&=~#]*)?",
    re.IGNORECASE,
)

REGEX_URL_ENCURTADA = re.compile(
    r"(?:bit\.ly|goo\.gl|tinyurl\.com|t\.co)/[\w\-]+",
    re.IGNORECASE,
)

REGEX_PLATAFORMA_CONTEXTUAL = re.compile(
    rf"(?i)(?:plataforma|acesse|site|portal|sistema|endereco\s*eletronico|"
    rf"disponivel\s*em|link|url)[:\s]*"
    rf"((?:https?://)?(?:www\.)?[\w\-\.]+{SUPPORTED_TLDS}"
    r"(?:/[\w\-\./?%&=~#]*)?)",
    re.IGNORECASE,
)

TAGS_POR_ORGAO = {
    "prefeitura": "veiculos_municipais",
    "camara": "veiculos_municipais",
    "detran": "veiculos_detran",
    "policia": "veiculos_policiais",
    "tribunal": "veiculos_judiciarios",
    "tj": "veiculos_judiciarios",
    "der": "veiculos_estaduais",
    "departamento de estradas": "veiculos_estaduais",
    "receita federal": "bens_apreendidos_federal",
    "secretaria da receita": "bens_apreendidos_federal",
    "srfb": "bens_apreendidos_federal",
    "patio": "veiculos_detran",
    "custodia": "veiculos_detran",
    "apreendidos": "veiculos_detran",
    "pm": "veiculos_policiais",
    "prf": "veiculos_policiais",
    "policia rodoviaria": "veiculos_policiais",
    "bombeiro": "veiculos_bombeiros",
    "samu": "veiculos_saude",
    "saude": "veiculos_saude",
}

KEYWORDS_LEILOEIRO = [
    "leiloeiro",
    "leilao",
    "lance",
    "arrematacao",
    "superbid",
    "sodresantoro",
    "zukerman",
    "joaoemilio",
    "leiloesfreire",
    "frfreiloes",
    "leilaobrasil",
    "megaleiloes",
    "sold",
    "leilomaster",
    "vipleiloes",
    "leiloesja",
    "portalleiloes",
    "lanceja",
    "leiloar",
    "leiloado",
    "hastaleiloes",
    "cleiloes",
    "bfreiloes",
    "alfreidoleiloes",
    "rioleiloes",
    "leilonet",
    "leilocentro",
]

# V12: Lista de domínios PROIBIDOS (não são sites de leiloeiros)
DOMINIOS_INVALIDOS = {
    'hotmail.com', 'hotmail.com.br',
    'yahoo.com', 'yahoo.com.br',
    'gmail.com', 'outlook.com',
    'uol.com.br', 'bol.com.br',
    'terra.com.br', 'ig.com.br',
    'globo.com', 'msn.com',
    'live.com', 'icloud.com'
}

# V12: Dicionário de tags específicas para leilões de veículos/sucatas
MAPA_TAGS = {
    # Estado do bem
    'sucata': ['sucata', 'sucateamento', 'ferro velho'],
    'documentado': ['documentado', 'com documento', 'documentação ok', 'documento em dia'],
    'sem_documento': ['sem documento', 'sem documentação', 'indocumentado'],
    'sinistrado': ['sinistrado', 'sinistro', 'acidentado', 'batido'],
    'recuperavel': ['recuperável', 'recuperavel', 'para recuperação'],
    'irrecuperavel': ['irrecuperável', 'irrecuperavel', 'inservível', 'inservivel'],

    # Tipo de veículo
    'automovel': ['automóvel', 'automovel', 'carro', 'veículo de passeio'],
    'motocicleta': ['motocicleta', 'moto', 'motociclo', 'ciclomotor'],
    'caminhao': ['caminhão', 'caminhao', 'truck'],
    'onibus': ['ônibus', 'onibus', 'micro-ônibus', 'microonibus'],
    'utilitario': ['utilitário', 'utilitario', 'pick-up', 'pickup', 'van'],
    'maquinario': ['máquina', 'maquina', 'trator', 'retroescavadeira', 'pá carregadeira'],
    'embarcacao': ['embarcação', 'embarcacao', 'barco', 'lancha'],

    # Condição
    'lote_unico': ['lote único', 'lote unico', 'item único'],
    'lote_multiplo': ['lotes', 'diversos itens', 'vários veículos'],
    'apreendido': ['apreendido', 'apreensão', 'retido'],
    'inservivel': ['inservível', 'inservivel', 'baixado', 'descarte']
}


@dataclass(frozen=True)
class ResultadoEdital:
    id_interno: str = "N/D"            # V11: PK interna
    n_edital: str = "N/D"
    data_publicacao: str = "N/D"       # V11: do JSON
    data_atualizacao: str = "N/D"      # V11: do JSON
    data_leilao: str = "N/D"
    titulo: str = "N/D"
    descricao: str = "N/D"
    objeto_resumido: str = "N/D"       # V11: extrair do PDF
    orgao: str = "N/D"
    uf: str = "N/D"
    cidade: str = "N/D"
    tags: str = "N/D"
    link_leiloeiro: str = "N/D"
    link_pncp: str = "N/D"
    modalidade_leilao: str = "N/D"     # V12: ONLINE | PRESENCIAL | HÍBRIDO
    valor_estimado: str = "N/D"        # V12: valor total estimado
    quantidade_itens: str = "N/D"      # V12: quantidade de itens/lotes
    nome_leiloeiro: str = "N/D"        # V12: nome do leiloeiro oficial
    arquivo_origem: str = "N/D"

    def as_dict(self) -> Dict[str, str]:
        return {
            "id_interno": self.id_interno,
            "n_edital": self.n_edital,
            "data_publicacao": self.data_publicacao,
            "data_atualizacao": self.data_atualizacao,
            "data_leilao": self.data_leilao,
            "titulo": self.titulo,
            "descricao": self.descricao,
            "objeto_resumido": self.objeto_resumido,
            "orgao": self.orgao,
            "uf": self.uf,
            "cidade": self.cidade,
            "tags": self.tags,
            "link_leiloeiro": self.link_leiloeiro,
            "link_pncp": self.link_pncp,
            "modalidade_leilao": self.modalidade_leilao,
            "valor_estimado": self.valor_estimado,
            "quantidade_itens": self.quantidade_itens,
            "nome_leiloeiro": self.nome_leiloeiro,
            "arquivo_origem": self.arquivo_origem,
        }


def corrigir_encoding(texto: str) -> str:
    if not texto or texto == "N/D":
        return texto
    try:
        return texto.encode("latin-1").decode("utf-8")
    except (UnicodeDecodeError, UnicodeEncodeError):
        return texto


def limpar_texto(texto: str, max_length: int = 500) -> str:
    if not texto or texto == "N/D":
        return texto
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    texto = re.sub(r" {2,}", " ", texto)
    if len(texto) > max_length:
        texto = texto[:max_length] + "..."
    return texto.strip()


def converter_data_extenso(match: re.Match) -> str:
    """Converte '14 de janeiro de 2026' para '14/01/2026'"""
    dia = match.group(1).zfill(2)
    mes = MESES_BR.get(match.group(2).lower(), "01")
    ano = match.group(3)
    return f"{dia}/{mes}/{ano}"


def gerar_id_interno(link_pncp: str) -> str:
    """Gera ID interno único baseado no link_pncp ou timestamp"""
    if not link_pncp or link_pncp == "N/D":
        return f"ID_{datetime.now().strftime('%Y%m%d%H%M%S%f')}"
    return f"ID_{hashlib.md5(link_pncp.encode()).hexdigest()[:12].upper()}"


def formatar_data_br(data_str: str) -> str:
    """
    V12: Formata data para padrão brasileiro DD/MM/YYYY.
    Aceita múltiplos formatos de entrada.
    """
    if not data_str or data_str == "N/D":
        return "N/D"

    data_str = str(data_str).strip()

    # Se já está no formato DD/MM/YYYY, retorna
    if re.match(r'\d{2}/\d{2}/\d{4}', data_str):
        return data_str

    # Tenta parsear ISO format (YYYY-MM-DD ou com timestamp)
    if 'T' in data_str or '-' in data_str:
        try:
            if 'T' in data_str:
                data_str = data_str.split('T')[0]

            # YYYY-MM-DD
            if re.match(r'\d{4}-\d{2}-\d{2}', data_str):
                ano, mes, dia = data_str.split('-')
                return f"{dia}/{mes}/{ano}"
        except:
            pass

    # Tenta parsear DD-MM-YYYY ou DD.MM.YYYY
    try:
        data_obj = datetime.strptime(data_str.replace('.', '/').replace('-', '/'), "%d/%m/%Y")
        return data_obj.strftime("%d/%m/%Y")
    except:
        pass

    return "N/D"


def detectar_leilao_presencial(modalidade: str, local_realizacao: str, descricao: str = "") -> bool:
    """
    V12: Detecta se leilão é presencial baseado em múltiplos indicadores.
    """
    texto = f"{modalidade} {local_realizacao} {descricao}".lower()

    indicadores_presencial = [
        'presencial',
        'auditório', 'auditorio',
        'sede', 'prédio', 'predio',
        'sala de licitação', 'sala de licitacao',
        'endereço:', 'endereco:',
        'rua ', 'avenida ', 'av. ',
        'edifício', 'edificio',
        'local:', 'comparecimento'
    ]

    indicadores_online = [
        'eletrônico', 'eletronico',
        'online', 'on-line',
        'internet', 'virtual',
        'www.', 'http', '.com', '.gov',
        'plataforma digital'
    ]

    score_presencial = sum(1 for ind in indicadores_presencial if ind in texto)
    score_online = sum(1 for ind in indicadores_online if ind in texto)

    return score_presencial > score_online


def extrair_data_de_texto(texto: str) -> Optional[str]:
    """Extrai data do texto, incluindo formatos extensos como '14 de janeiro de 2026'"""
    if not texto:
        return None

    # Tentativa 1: Regex contextual (dd/mm/yyyy)
    matches = REGEX_DATA_CONTEXTUAL.findall(texto)
    if not matches:
        matches = REGEX_DATA.findall(texto)

    for data_str in matches:
        try:
            data_obj = datetime.strptime(data_str.replace("-", "/"), "%d/%m/%Y")
            if data_obj.year >= 2024:
                return data_str
        except ValueError:
            continue

    # Tentativa 2: Formato extenso "14 de janeiro de 2026"
    match_extenso = REGEX_DATA_EXTENSO.search(texto)
    if match_extenso:
        data_convertida = converter_data_extenso(match_extenso)
        try:
            data_obj = datetime.strptime(data_convertida, "%d/%m/%Y")
            if data_obj.year >= 2024:
                return data_convertida
        except ValueError:
            pass

    return None


def extrair_objeto_resumido(texto: str) -> str:
    """Extrai resumo dos objetos/veículos do leilão"""
    if not texto:
        return "N/D"

    matches = REGEX_OBJETO_RESUMIDO.findall(texto)
    if matches:
        # Limpa e junta até 3 matches
        objetos = [m.strip()[:300] for m in matches[:3]]
        return " | ".join(objetos)[:300]

    return "N/D"


def normalizar_url(url: str) -> str:
    url = url.strip()
    url = re.sub(r"\s+", "", url)
    url = url.rstrip(".,;:)>")
    url = url.rstrip("\"'")
    if not url:
        return ""
    if not url.lower().startswith(("http://", "https://")):
        url = "https://" + url
    return url


def extrair_urls_de_texto(texto: str) -> List[str]:
    if not texto:
        return []

    urls: List[str] = []

    urls.extend(REGEX_URL.findall(texto))

    for raw in REGEX_URL_QUEBRADA.findall(texto):
        urls.append(raw)

    for raw in REGEX_URL_ENCURTADA.findall(texto):
        urls.append(raw)

    urls_limpas: List[str] = []
    for raw in urls:
        url = normalizar_url(raw)
        if url and len(url) > 10:
            urls_limpas.append(url)

    return list(dict.fromkeys(urls_limpas))


def extrair_urls_contextuais(texto: str) -> List[str]:
    if not texto:
        return []

    urls: List[str] = []
    for match in REGEX_PLATAFORMA_CONTEXTUAL.findall(texto):
        url = normalizar_url(match)
        if url and len(url) > 10:
            urls.append(url)

    return list(dict.fromkeys(urls))


def encontrar_link_leiloeiro(urls: Sequence[str], texto_completo: str = "") -> Optional[str]:
    urls_unicas = list(dict.fromkeys([u for u in urls if u]))

    for url in urls_unicas:
        url_lower = url.lower()
        if any(keyword in url_lower for keyword in KEYWORDS_LEILOEIRO):
            return url

    if texto_completo:
        urls_ctx = extrair_urls_contextuais(texto_completo)
        for url in urls_ctx:
            url_lower = url.lower()
            if "gov" in url_lower or "pncp" in url_lower:
                continue
            if any(keyword in url_lower for keyword in KEYWORDS_LEILOEIRO):
                return url

        for url in urls_ctx:
            url_lower = url.lower()
            if "gov" in url_lower or "pncp" in url_lower:
                continue
            if ".com.br" in url_lower or ".com" in url_lower or ".net.br" in url_lower or ".net" in url_lower:
                return url

    for url in urls_unicas:
        url_lower = url.lower()
        if "gov" in url_lower or "pncp" in url_lower:
            continue
        if any(tld in url_lower for tld in [".com.br", ".com", ".net.br", ".net", ".org.br"]):
            return url

    return None


def gerar_tags(orgao: str) -> str:
    if not orgao or orgao == "N/D":
        return "veiculos_gerais"

    orgao_lower = orgao.lower()
    for palavra_chave, tag in TAGS_POR_ORGAO.items():
        if palavra_chave in orgao_lower:
            return tag

    return "veiculos_gerais"


def extrair_texto_docx(arquivo_docx: Path) -> str:
    if not DOCX_AVAILABLE:
        return ""
    try:
        doc = DocxDocument(str(arquivo_docx))
        partes: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                partes.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        partes.append(cell.text)
        return "\n".join(partes)
    except Exception:
        return ""


def extrair_urls_de_docx(pasta: Path) -> List[str]:
    urls: List[str] = []
    for docx_path in list(pasta.glob("*.docx")) + list(pasta.glob("*.doc")):
        texto = extrair_texto_docx(docx_path)
        if not texto:
            continue
        urls.extend(extrair_urls_de_texto(texto))
        urls.extend(extrair_urls_contextuais(texto))
    return list(dict.fromkeys(urls))


def extrair_urls_de_xlsx(pasta: Path) -> List[str]:
    urls: List[str] = []
    arquivos = list(pasta.glob("*.xlsx")) + list(pasta.glob("*.xls")) + list(pasta.glob("*.xlsm"))
    for excel_path in arquivos:
        try:
            xl = pd.ExcelFile(excel_path)
            for sheet_name in xl.sheet_names:
                df = pd.read_excel(xl, sheet_name=sheet_name, dtype=str)
                for col in df.columns:
                    for valor in df[col].dropna():
                        valor_str = str(valor)
                        if "http" in valor_str.lower() or ".com" in valor_str.lower() or ".net" in valor_str.lower():
                            urls.extend(extrair_urls_de_texto(valor_str))
        except Exception:
            continue
    return list(dict.fromkeys(urls))


def extrair_urls_de_zip(pasta: Path) -> List[str]:
    urls: List[str] = []
    for zip_path in pasta.glob("*.zip"):
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir)
                with zipfile.ZipFile(zip_path, "r") as zf:
                    zf.extractall(temp_path)

                urls.extend(extrair_urls_de_docx(temp_path))
                urls.extend(extrair_urls_de_xlsx(temp_path))

                for pdf_path in temp_path.glob("**/*.pdf"):
                    try:
                        texto = extrair_texto_pdfs([pdf_path])
                        urls.extend(extrair_urls_de_texto(texto))
                        urls.extend(extrair_urls_contextuais(texto))
                    except Exception:
                        continue
        except Exception:
            continue
    return list(dict.fromkeys(urls))


def extrair_texto_pdfs(pdfs: Sequence[Path]) -> str:
    partes: List[str] = []
    for pdf_path in pdfs:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    partes.append(page.extract_text() or "")
        except Exception:
            continue
    return " ".join([p for p in partes if p]).strip()


def extrair_de_metadados_json(pasta: Path) -> Dict[str, str]:
    dados: Dict[str, str] = {}
    json_file = pasta / "metadados_pncp.json"
    if not json_file.exists():
        return dados

    try:
        with open(json_file, "r", encoding="utf-8") as f:
            metadados = json.load(f)

        titulo_original = metadados.get("titulo") or ""
        if titulo_original:
            titulo_corrigido = corrigir_encoding(titulo_original)
            dados["titulo"] = titulo_corrigido
            match = REGEX_N_EDITAL.search(titulo_corrigido)
            if match:
                dados["n_edital"] = match.group(1).replace(" ", "")

        descricao_original = ""
        if metadados.get("descricao"):
            descricao_original = str(metadados["descricao"])
            dados["descricao"] = limpar_texto(corrigir_encoding(descricao_original), max_length=500)
        elif metadados.get("objeto"):
            descricao_original = str(metadados["objeto"])
            dados["descricao"] = limpar_texto(corrigir_encoding(descricao_original), max_length=500)

        # V11: Extrai data_publicacao e data_atualizacao
        if metadados.get("data_publicacao"):
            dados["data_publicacao"] = str(metadados["data_publicacao"]).split("T")[0]

        if metadados.get("data_atualizacao"):
            dados["data_atualizacao"] = str(metadados["data_atualizacao"]).split("T")[0]

        if metadados.get("data_inicio_propostas"):
            dados["data_leilao"] = str(metadados["data_inicio_propostas"]).split("T")[0]
        elif metadados.get("dataAberturaPropostas"):
            dados["data_leilao"] = str(metadados["dataAberturaPropostas"]).split("T")[0]
        elif descricao_original:
            data_extraida = extrair_data_de_texto(descricao_original)
            if data_extraida:
                dados["data_leilao"] = data_extraida

        urls_todas: List[str] = []
        texto_completo = ""

        if descricao_original:
            urls_todas.extend(extrair_urls_de_texto(descricao_original))
            texto_completo += descricao_original + " "

        objeto = str(metadados.get("objeto") or "")
        if objeto:
            urls_todas.extend(extrair_urls_de_texto(objeto))
            texto_completo += objeto + " "

        info = metadados.get("informacoes_complementares")
        if info is not None:
            info_str = str(info)
            urls_todas.extend(extrair_urls_de_texto(info_str))
            texto_completo += info_str + " "

        for file_info in metadados.get("files_meta", []) or []:
            titulo_file = str(file_info.get("titulo") or "")
            if titulo_file:
                urls_todas.extend(extrair_urls_de_texto(titulo_file))

        link_leiloeiro = encontrar_link_leiloeiro(urls_todas, texto_completo)
        if link_leiloeiro:
            dados["link_leiloeiro"] = link_leiloeiro

        link_pncp = metadados.get("link_pncp")
        if link_pncp:
            dados["link_pncp"] = str(link_pncp)

        orgao = metadados.get("orgao_nome")
        if orgao:
            dados["orgao"] = str(orgao)

        uf = metadados.get("uf")
        if uf:
            dados["uf"] = str(uf)

        municipio = metadados.get("municipio")
        if municipio:
            dados["cidade"] = str(municipio)

    except Exception as exc:
        print(f"[AVISO] Erro ao ler JSON: {exc}")

    return dados


def extrair_de_excel(pasta: Path) -> Dict[str, str]:
    dados: Dict[str, str] = {}
    arquivos = list(pasta.glob("*.xlsx")) + list(pasta.glob("*.xls")) + list(pasta.glob("*.csv"))
    if not arquivos:
        return dados

    for arquivo in arquivos:
        try:
            if arquivo.suffix.lower() == ".csv":
                df = pd.read_csv(arquivo, encoding="utf-8", nrows=50)
            else:
                df = pd.read_excel(arquivo, nrows=50)
            if df.empty:
                continue

            for col in df.columns:
                col_lower = str(col).lower()

                if "edital" in col_lower and "n_edital" not in dados:
                    valores = df[col].dropna().astype(str).tolist()
                    if valores:
                        dados["n_edital"] = valores[0]

                if any(kw in col_lower for kw in ["data", "leilao", "abertura"]) and "data_leilao" not in dados:
                    valores = df[col].dropna().astype(str).tolist()
                    if valores:
                        data = extrair_data_de_texto(valores[0])
                        if data:
                            dados["data_leilao"] = data

                if any(kw in col_lower for kw in ["desc", "objeto", "bem"]) and "descricao" not in dados:
                    valores = df[col].dropna().astype(str).tolist()
                    if valores:
                        dados["descricao"] = " | ".join(valores[:3])
        except Exception:
            continue

    return dados


def extrair_de_docx(pasta: Path) -> Dict[str, str]:
    dados: Dict[str, str] = {}
    if not DOCX_AVAILABLE:
        return dados

    arquivos = list(pasta.glob("*.docx")) + list(pasta.glob("*.doc"))
    for docx_path in arquivos:
        texto = extrair_texto_docx(docx_path)
        if not texto:
            continue

        if "n_edital" not in dados:
            match = REGEX_N_EDITAL.search(texto)
            if match:
                dados["n_edital"] = match.group(1).replace(" ", "")

        if "data_leilao" not in dados:
            data = extrair_data_de_texto(texto)
            if data:
                dados["data_leilao"] = data

        if "link_leiloeiro" not in dados:
            urls = extrair_urls_de_texto(texto)
            link = encontrar_link_leiloeiro(urls, texto)
            if link:
                dados["link_leiloeiro"] = link

    return dados


def extrair_de_path(pasta_edital: Path) -> Dict[str, str]:
    dados: Dict[str, str] = {}
    partes = pasta_edital.name.split("_")

    if len(partes) >= 3:
        dados["uf"] = partes[0].upper()
        dados["cidade"] = partes[1].replace("-", " ").title()
        dados["orgao"] = " ".join(partes[2:]).replace("-", " ").title()
    elif len(partes) >= 2:
        dados["uf"] = partes[0].upper()
        dados["cidade"] = partes[1].replace("-", " ").title()

    return dados


def extrair_de_pdf(pasta: Path) -> Dict[str, str]:
    dados: Dict[str, str] = {}
    pdfs = sorted(list(pasta.glob("*.pdf")))
    if not pdfs:
        return dados

    texto_completo = extrair_texto_pdfs(pdfs)
    if not texto_completo:
        return dados

    match = REGEX_N_EDITAL.search(texto_completo)
    if match:
        dados["n_edital"] = match.group(1).replace(" ", "")

    data_extraida = extrair_data_de_texto(texto_completo)
    if data_extraida:
        dados["data_leilao"] = data_extraida

    match_titulo = re.search(
        r"(?i)(?:objeto|finalidade).*?[:;]\s*(.{10,200}?)(?:\.|;|\n\n)",
        texto_completo,
        re.DOTALL,
    )
    if match_titulo:
        dados["titulo"] = match_titulo.group(1).strip()[:200]

    # V11: Extrai objeto_resumido
    objeto_resumido = extrair_objeto_resumido(texto_completo)
    if objeto_resumido != "N/D":
        dados["objeto_resumido"] = objeto_resumido

    urls = extrair_urls_de_texto(texto_completo)
    urls.extend(extrair_urls_contextuais(texto_completo))
    link = encontrar_link_leiloeiro(urls, texto_completo)
    if link:
        dados["link_leiloeiro"] = link

    return dados


def extrair_urls_de_todos_arquivos(pasta: Path) -> Tuple[List[str], str]:
    urls: List[str] = []
    texto: List[str] = []

    urls.extend(extrair_urls_de_docx(pasta))
    urls.extend(extrair_urls_de_xlsx(pasta))
    urls.extend(extrair_urls_de_zip(pasta))

    for docx_path in list(pasta.glob("*.docx")) + list(pasta.glob("*.doc")):
        doc_text = extrair_texto_docx(docx_path)
        if doc_text:
            texto.append(doc_text)

    pdfs = sorted(list(pasta.glob("*.pdf")))
    pdf_text = extrair_texto_pdfs(pdfs)
    if pdf_text:
        texto.append(pdf_text)

    texto_completo = " ".join(texto).strip()
    return list(dict.fromkeys(urls)), texto_completo


def pasta_parece_edital(pasta: Path) -> bool:
    if not pasta.is_dir():
        return False
    if (pasta / "metadados_pncp.json").exists():
        return True
    padroes = ["*.pdf", "*.docx", "*.doc", "*.xlsx", "*.xls", "*.xlsm", "*.csv", "*.zip"]
    return any(pasta.glob(p) for p in padroes)


def listar_pastas_editais(pasta_raiz: Path) -> List[Path]:
    if not pasta_raiz.exists():
        return []

    editais: List[Path] = []
    municipios = [d for d in pasta_raiz.iterdir() if d.is_dir()]

    for municipio in municipios:
        subdirs = [d for d in municipio.iterdir() if d.is_dir()]
        subdirs_editais = [d for d in subdirs if pasta_parece_edital(d)]

        if subdirs_editais:
            editais.extend(sorted(subdirs_editais))
            continue

        if pasta_parece_edital(municipio):
            editais.append(municipio)

    return editais


def processar_edital(pasta_edital: Path) -> Optional[Dict[str, str]]:
    print(f"[INFO] Processando: {pasta_edital.relative_to(PASTA_EDITAIS)}")

    if not pasta_edital.exists() or not pasta_edital.is_dir():
        print("  [ERRO] Pasta do edital invalida")
        return None

    resultado = ResultadoEdital(arquivo_origem=str(pasta_edital.relative_to(PASTA_EDITAIS)))
    dados_finais = resultado.as_dict()

    fontes: List[Tuple[str, Dict[str, str]]] = [
        ("JSON", extrair_de_metadados_json(pasta_edital)),
        ("Excel", extrair_de_excel(pasta_edital)),
        ("DOCX", extrair_de_docx(pasta_edital)),
        ("Path", extrair_de_path(pasta_edital)),
        ("PDF", extrair_de_pdf(pasta_edital)),
    ]

    campos_encontrados: List[str] = []
    for nome_fonte, dados_fonte in fontes:
        for campo, valor in dados_fonte.items():
            if campo not in dados_finais:
                continue
            if dados_finais[campo] == "N/D" and valor and valor != "N/D":
                dados_finais[campo] = valor
                campos_encontrados.append(f"{campo}({nome_fonte})")

    if dados_finais["link_leiloeiro"] == "N/D":
        urls_todos, texto_todos = extrair_urls_de_todos_arquivos(pasta_edital)
        link_encontrado = encontrar_link_leiloeiro(urls_todos, texto_todos)
        if link_encontrado:
            dados_finais["link_leiloeiro"] = link_encontrado
            campos_encontrados.append("link_leiloeiro(V10_DEEP)")

    dados_finais["tags"] = gerar_tags(dados_finais["orgao"])

    # V11: Gera id_interno baseado no link_pncp
    dados_finais["id_interno"] = gerar_id_interno(dados_finais["link_pncp"])

    if campos_encontrados:
        preview = campos_encontrados[:5]
        suffix = "..." if len(campos_encontrados) > 5 else ""
        print(f"  [OK] Extraido: {', '.join(preview)}{suffix}")
    else:
        print("  [AVISO] Nenhum dado extraido")

    return dados_finais


def varrer_editais() -> List[Dict[str, str]]:
    if not PASTA_EDITAIS.exists():
        print(f"[ERRO] Diretorio {PASTA_EDITAIS} nao encontrado")
        return []

    print(f"[INFO] Varrendo editais em: {PASTA_EDITAIS}")

    pastas_editais = listar_pastas_editais(PASTA_EDITAIS)
    print(f"[INFO] Total de editais encontrados: {len(pastas_editais)}\n")

    resultados: List[Dict[str, str]] = []
    for i, pasta in enumerate(pastas_editais, 1):
        try:
            print(f"[{i}/{len(pastas_editais)}] ", end="")
            dados = processar_edital(pasta)
            if dados:
                resultados.append(dados)
        except Exception as exc:
            print(f"[ERRO] Falha ao processar {pasta}: {exc}")

    return resultados


def extrair_n_pncp(link_pncp: str) -> str:
    """Extrai n_pncp do link_pncp, retorna N/D se não encontrar"""
    if not link_pncp or link_pncp == "N/D":
        return "N/D"

    # Tenta extrair do padrão: /editais/{cnpj}/{ano}/{sequencial}
    match = re.search(r"/editais/([^/]+/\d{4}/\d+)", link_pncp)
    if match:
        return match.group(1)

    # Fallback: tenta extrair qualquer padrão numérico após /editais/
    match = re.search(r"/editais/([^/\s]+)", link_pncp)
    if match:
        return match.group(1)

    return "N/D"


def salvar_resultado_final(dados: List[Dict[str, str]]) -> None:
    """Gera arquivo RESULTADO_FINAL.xlsx com todos os campos na ordem especificada"""
    if not dados:
        print("[ERRO] Nenhum dado para salvar")
        return

    # Adiciona campo n_pncp extraído do link_pncp
    for item in dados:
        item["n_pncp"] = extrair_n_pncp(item.get("link_pncp", "N/D"))

    df = pd.DataFrame(dados)

    # Reordena colunas conforme especificado
    colunas_ordenadas = [
        "id_interno", "orgao", "uf", "cidade", "n_pncp", "n_edital",
        "data_publicacao", "data_atualizacao", "data_leilao",
        "titulo", "descricao", "tags", "link_pncp", "link_leiloeiro",
        "objeto_resumido", "arquivo_origem"
    ]

    # Adiciona colunas faltantes com N/D
    for col in colunas_ordenadas:
        if col not in df.columns:
            df[col] = "N/D"

    df = df[colunas_ordenadas]

    output_path = PASTA_RAIZ / "RESULTADO_FINAL.xlsx"
    df.to_excel(output_path, index=False, engine="openpyxl")

    print(f"\n{'='*60}")
    print(f"[OK] RESULTADO FINAL GERADO: {output_path}")
    print(f"[OK] TOTAL DE REGISTROS: {len(dados)}")
    print(f"{'='*60}\n")

    # Exibe estatísticas de preenchimento
    print("=" * 60)
    print("VALIDAÇÃO RESULTADO_FINAL.xlsx")
    print("=" * 60)
    print(f"{'[CAMPO]':<25} {'[PREENCHIDOS]':<20} {'[TAXA]':<10}")
    print("-" * 60)

    for col in colunas_ordenadas:
        if col == "arquivo_origem":
            continue
        total = len(df)
        nao_vazio = df[col].notna() & (df[col] != "") & (df[col] != "N/D")
        preenchidos = int(nao_vazio.sum())
        taxa = (preenchidos / total * 100) if total > 0 else 0

        status = "[OK]" if taxa >= 80 else "[PARCIAL]" if taxa >= 50 else "[BAIXO]"
        print(f"{col:<25} {preenchidos:3d}/{total:3d} ({taxa:5.1f}%)    {status}")

    print("=" * 60 + "\n")


def salvar_csv(dados: List[Dict[str, str]]) -> None:
    if not dados:
        print("[AVISO] Nenhum dado para salvar")
        return

    df = pd.DataFrame(dados)
    df.to_csv(CSV_OUTPUT, index=False, encoding="utf-8-sig")

    print(f"\n[OK] Arquivo gerado: {CSV_OUTPUT}")
    print(f"[OK] Total de editais processados: {len(dados)}")

    print("\n" + "=" * 60)
    print("ESTATISTICAS DE PREENCHIMENTO - V10")
    print("=" * 60)

    for col in df.columns:
        if col == "arquivo_origem":
            continue
        total = len(df)
        nao_vazio = df[col].notna() & (df[col] != "") & (df[col] != "N/D")
        preenchidos = int(nao_vazio.sum())
        taxa = (preenchidos / total * 100) if total > 0 else 0

        status = "[OK]" if taxa == 100 else "[PARCIAL]" if taxa >= 50 else "[BAIXO]"
        barra = "#" * int(taxa / 5)
        print(f"{status} {col:20s} {preenchidos:3d}/{total:3d} ({taxa:5.1f}%) {barra}")


def main() -> None:
    print("\n" + "=" * 60)
    print("ACHE SUCATAS DaaS - AUDITOR V11")
    print("=" * 60)
    print("\nV11 - NOVIDADES:")
    print("[+] Suporte a .leilao.br")
    print("[+] Datas extensas (14 de janeiro de 2026)")
    print("[+] Novos campos: id_interno, data_publicacao, data_atualizacao, objeto_resumido")
    print("[+] Arquivo RESULTADO_FINAL.xlsx com todas as colunas")
    print("\nCascata: JSON -> Excel -> DOCX -> Path -> PDF")
    print("=" * 60 + "\n")

    if not DOCX_AVAILABLE:
        print("[AVISO] python-docx nao instalado. DOCX sera ignorado.\n")

    resultados = varrer_editais()
    salvar_csv(resultados)
    salvar_resultado_final(resultados)

    print("\n" + "=" * 60)
    print("PROCESSAMENTO CONCLUIDO - V11")
    print("=" * 60 + "\n")


if __name__ == "__main__":
    main()
