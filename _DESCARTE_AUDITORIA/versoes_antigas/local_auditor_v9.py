#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ACHE SUCATAS DaaS - AUDITOR V9
Extrai dados de editais usando cascata: JSON -> Excel -> DOCX -> Path -> PDF

Melhorias V9:
- Fallback hierarquico para link_leiloeiro (plataforma, acesse, portal)
- Regex de URLs expandido (quebras de linha, parenteses, encurtadores)
- Keywords de leiloeiro expandidas (+6 plataformas)
- Extracao contextual em PDFs

Versao: 9.0
"""

import re
import json
import zipfile
import tempfile
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional

import pdfplumber
import pandas as pd

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[AVISO] python-docx nao instalado. Instale com: pip install python-docx")


# ============================================================================
# CONFIGURACOES
# ============================================================================

PASTA_RAIZ = Path(__file__).parent
PASTA_EDITAIS = PASTA_RAIZ / "ACHE_SUCATAS_DB"
CSV_OUTPUT = PASTA_RAIZ / "analise_editais_v9.csv"


# ============================================================================
# REGEX PATTERNS - V9 EXPANDIDO
# ============================================================================

REGEX_N_EDITAL = re.compile(
    r"(?i)(?:edital|processo|leilao|pregao).*?(\d{1,5}\s*/\s*20\d{2})",
    re.DOTALL
)

REGEX_DATA = re.compile(r"\b(\d{1,2}[/\-]\d{1,2}[/\-]20\d{2})\b")

# V9: Regex de URL expandido para capturar mais padroes
REGEX_URL = re.compile(
    r"https?://(?:www\.)?[\w\-\.]+(?:\.com\.br|\.com|\.org\.br|\.gov\.br|\.net)"
    r"(?:/[\w\-\./?%&=~#]*)?",
    re.IGNORECASE
)

# V9: Regex para URLs com quebra de linha em PDFs
REGEX_URL_QUEBRADA = re.compile(
    r"https?://(?:www\.)?[\w\-\.]+\s*[\n\r]\s*(?:\.com\.br|\.com|\.org\.br)"
    r"(?:/[\w\-\./?%&=]*)?",
    re.IGNORECASE
)

# V9: Regex para URLs encurtadas
REGEX_URL_ENCURTADA = re.compile(
    r"(?:bit\.ly|goo\.gl|tinyurl\.com|t\.co)/[\w\-]+",
    re.IGNORECASE
)

# V9: Regex contextual para extracao de plataformas
REGEX_PLATAFORMA_CONTEXTUAL = re.compile(
    r"(?i)(?:plataforma|acesse|site|portal|sistema|endereco\s*eletronico|"
    r"disponivel\s*em|link|url)[:\s]*"
    r"((?:https?://)?(?:www\.)?[\w\-]+\.(?:com\.br|com|org\.br)(?:/[\w\-\./?%&=]*)?)",
    re.IGNORECASE
)

REGEX_DATA_CONTEXTUAL = re.compile(
    r"(?i)(?:data|abertura|sessao|leilao|pregao|realizacao|"
    r"desfazimento|alienacao|hasta|arrematacao).*?"
    r"(\d{1,2}[/\-]\d{1,2}[/\-]20\d{2})",
    re.DOTALL
)


# ============================================================================
# DICIONARIO DE TAGS - V9
# ============================================================================

TAGS_POR_ORGAO = {
    'prefeitura': 'veiculos_municipais',
    'camara': 'veiculos_municipais',
    'detran': 'veiculos_detran',
    'policia': 'veiculos_policiais',
    'tribunal': 'veiculos_judiciarios',
    'tj': 'veiculos_judiciarios',
    'der': 'veiculos_estaduais',
    'departamento de estradas': 'veiculos_estaduais',
    'receita federal': 'bens_apreendidos_federal',
    'secretaria da receita': 'bens_apreendidos_federal',
    'srfb': 'bens_apreendidos_federal',
    'patio': 'veiculos_detran',
    'custodia': 'veiculos_detran',
    'apreendidos': 'veiculos_detran',
    'pm': 'veiculos_policiais',
    'prf': 'veiculos_policiais',
    'policia rodoviaria': 'veiculos_policiais',
    'bombeiro': 'veiculos_bombeiros',
    'samu': 'veiculos_saude',
    'saude': 'veiculos_saude',
}


# ============================================================================
# KEYWORDS DE LEILOEIRO - V9 EXPANDIDO
# ============================================================================

KEYWORDS_LEILOEIRO = [
    # Plataformas consolidadas
    'leiloeiro',
    'leilao',
    'lance',
    'arrematacao',
    'superbid',
    'sodresantoro',
    'zukerman',
    'joaoemilio',
    'leiloesfreire',
    'frfreiloes',
    'leilaobrasil',
    'megaleiloes',
    'sold',
    'leilomaster',
    'vipleiloes',
    'leiloesja',
    'portalleiloes',
    # V9: Novas plataformas identificadas
    'lanceja',
    'leiloar',
    'leiloado',
    'hastaleiloes',
    'cleiloes',
    'bfreiloes',
    'alfreidoleiloes',
    'rioleiloes',
    'leilonet',
    'leilocentro',
]

# V9: Termos contextuais para fallback
TERMOS_CONTEXTO_PLATAFORMA = [
    'plataforma',
    'acesse',
    'acessar',
    'portal',
    'site',
    'sistema',
    'endereco eletronico',
    'disponivel em',
    'lances em',
    'lances no',
    'cadastro em',
    'cadastre-se',
]


# ============================================================================
# FUNCOES AUXILIARES
# ============================================================================

def corrigir_encoding(texto: str) -> str:
    """Corrige dupla codificacao UTF-8."""
    if not texto or texto == 'N/D':
        return texto
    try:
        return texto.encode('latin-1').decode('utf-8')
    except (UnicodeDecodeError, UnicodeEncodeError):
        return texto


def limpar_texto(texto: str, max_length: int = 500) -> str:
    """Limpa e trunca texto."""
    if not texto or texto == 'N/D':
        return texto

    texto = re.sub(r'\n{3,}', '\n\n', texto)
    texto = re.sub(r' {2,}', ' ', texto)

    if len(texto) > max_length:
        texto = texto[:max_length] + '...'

    return texto.strip()


def extrair_data_de_texto(texto: str) -> Optional[str]:
    """Extrai data do leilao de um texto usando contexto."""
    if not texto:
        return None

    matches = REGEX_DATA_CONTEXTUAL.findall(texto)
    if not matches:
        matches = REGEX_DATA.findall(texto)

    for data_str in matches:
        try:
            data_obj = datetime.strptime(data_str.replace('-', '/'), '%d/%m/%Y')
            if data_obj.year >= 2024:
                return data_str
        except ValueError:
            continue

    return None


def extrair_urls_de_texto(texto: str) -> List[str]:
    """Extrai todas as URLs de um texto usando multiplos padroes."""
    if not texto:
        return []

    urls = []

    # Padrao principal
    urls.extend(REGEX_URL.findall(texto))

    # URLs com quebra de linha
    urls_quebradas = REGEX_URL_QUEBRADA.findall(texto)
    for url in urls_quebradas:
        url_limpa = re.sub(r'\s+', '', url)
        urls.append(url_limpa)

    # URLs encurtadas
    urls_encurtadas = REGEX_URL_ENCURTADA.findall(texto)
    for url in urls_encurtadas:
        if not url.startswith('http'):
            url = 'https://' + url
        urls.append(url)

    # Limpar URLs
    urls_limpas = []
    for url in urls:
        url = url.rstrip('.,;:)>')
        url = url.rstrip('"\'')
        if url and len(url) > 10:
            urls_limpas.append(url)

    return urls_limpas


def extrair_urls_contextuais(texto: str) -> List[str]:
    """V9: Extrai URLs usando contexto (plataforma, acesse, portal)."""
    if not texto:
        return []

    urls = []

    matches = REGEX_PLATAFORMA_CONTEXTUAL.findall(texto)
    for match in matches:
        url = match.strip()
        if not url.startswith('http'):
            url = 'https://' + url
        url = url.rstrip('.,;:)>')
        if url and len(url) > 10:
            urls.append(url)

    return urls


def encontrar_link_leiloeiro(urls: List[str], texto_completo: str = "") -> Optional[str]:
    """
    Encontra o link do leiloeiro usando cascata hierarquica.
    
    V9: Fallback hierarquico completo:
    1. Keywords especificas de leiloeiro na URL
    2. Extracao contextual (plataforma, acesse, portal)
    3. Dominios .com.br que nao sejam gov/pncp
    """
    if not urls:
        urls = []

    urls_unicas = list(dict.fromkeys(urls))

    # Nivel 1: Keywords especificas de leiloeiro
    for url in urls_unicas:
        url_lower = url.lower()
        for keyword in KEYWORDS_LEILOEIRO:
            if keyword in url_lower:
                return url

    # Nivel 2: Extracao contextual do texto completo
    if texto_completo:
        urls_contextuais = extrair_urls_contextuais(texto_completo)
        for url in urls_contextuais:
            url_lower = url.lower()
            if 'gov' not in url_lower and 'pncp' not in url_lower:
                for keyword in KEYWORDS_LEILOEIRO:
                    if keyword in url_lower:
                        return url

        for url in urls_contextuais:
            url_lower = url.lower()
            if 'gov' not in url_lower and 'pncp' not in url_lower:
                if '.com.br' in url_lower or '.com' in url_lower:
                    return url

    # Nivel 3: Dominios .com.br genericos (nao gov/pncp)
    for url in urls_unicas:
        url_lower = url.lower()
        if '.com' in url_lower:
            if 'gov' not in url_lower and 'pncp' not in url_lower:
                return url

    return None


def encontrar_pasta_dados(pasta_edital: Path) -> Optional[Path]:
    """Encontra a pasta real com os dados."""
    subpastas = [d for d in pasta_edital.iterdir() if d.is_dir()]
    if subpastas:
        return subpastas[0]
    return pasta_edital


# ============================================================================
# EXTRACAO DE DOCX
# ============================================================================

def extrair_texto_docx(arquivo_docx: Path) -> str:
    """Extrai texto completo de um arquivo DOCX."""
    if not DOCX_AVAILABLE:
        return ""

    try:
        doc = DocxDocument(str(arquivo_docx))
        texto_completo = []

        for para in doc.paragraphs:
            if para.text.strip():
                texto_completo.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texto_completo.append(cell.text)

        return '\n'.join(texto_completo)

    except Exception:
        return ""


def extrair_urls_de_docx(pasta_dados: Path) -> List[str]:
    """Extrai URLs de todos os arquivos DOCX na pasta."""
    urls_encontradas = []
    arquivos_docx = list(pasta_dados.glob("*.docx")) + list(pasta_dados.glob("*.doc"))

    for docx_path in arquivos_docx:
        texto = extrair_texto_docx(docx_path)
        if texto:
            urls = extrair_urls_de_texto(texto)
            urls_encontradas.extend(urls)
            urls_ctx = extrair_urls_contextuais(texto)
            urls_encontradas.extend(urls_ctx)

    return urls_encontradas


# ============================================================================
# EXTRACAO DE XLSX
# ============================================================================

def extrair_urls_de_xlsx(pasta_dados: Path) -> List[str]:
    """Extrai URLs de todos os arquivos Excel na pasta."""
    urls_encontradas = []

    arquivos_excel = (
        list(pasta_dados.glob("*.xlsx")) +
        list(pasta_dados.glob("*.xls")) +
        list(pasta_dados.glob("*.xlsm"))
    )

    for excel_path in arquivos_excel:
        try:
            xl = pd.ExcelFile(excel_path)
            for sheet_name in xl.sheet_names:
                df = pd.read_excel(xl, sheet_name=sheet_name, dtype=str)
                for col in df.columns:
                    for valor in df[col].dropna():
                        valor_str = str(valor)
                        if 'http' in valor_str.lower() or '.com' in valor_str.lower():
                            urls = extrair_urls_de_texto(valor_str)
                            urls_encontradas.extend(urls)
        except Exception:
            pass

    return urls_encontradas


# ============================================================================
# PROCESSAMENTO DE ZIP
# ============================================================================

def extrair_urls_de_zip(pasta_dados: Path) -> List[str]:
    """Extrai URLs de arquivos dentro de ZIPs."""
    urls_encontradas = []
    arquivos_zip = list(pasta_dados.glob("*.zip"))

    for zip_path in arquivos_zip:
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir)

                with zipfile.ZipFile(zip_path, 'r') as zf:
                    zf.extractall(temp_path)

                urls_encontradas.extend(extrair_urls_de_docx(temp_path))
                urls_encontradas.extend(extrair_urls_de_xlsx(temp_path))

                for pdf_path in temp_path.glob("**/*.pdf"):
                    try:
                        with pdfplumber.open(pdf_path) as pdf:
                            for page in pdf.pages:
                                texto = page.extract_text() or ""
                                urls = extrair_urls_de_texto(texto)
                                urls_encontradas.extend(urls)
                                urls_ctx = extrair_urls_contextuais(texto)
                                urls_encontradas.extend(urls_ctx)
                    except Exception:
                        pass

        except Exception:
            pass

    return urls_encontradas


# ============================================================================
# FUNCOES DE EXTRACAO - CASCATA DE FONTES
# ============================================================================

def extrair_de_metadados_json(pasta_dados: Path) -> Dict[str, str]:
    """Prioridade 1: Extrai dados do arquivo metadados_pncp.json."""
    dados = {}
    json_file = pasta_dados / "metadados_pncp.json"

    if not json_file.exists():
        return dados

    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            metadados = json.load(f)

        if metadados.get('titulo'):
            titulo_corrigido = corrigir_encoding(metadados['titulo'])
            dados['titulo'] = titulo_corrigido
            match = REGEX_N_EDITAL.search(titulo_corrigido)
            if match:
                dados['n_edital'] = match.group(1).replace(' ', '')

        descricao_original = ''
        if metadados.get('descricao'):
            descricao_original = metadados['descricao']
            descricao_corrigida = corrigir_encoding(descricao_original)
            dados['descricao'] = limpar_texto(descricao_corrigida, max_length=500)
        elif metadados.get('objeto'):
            descricao_original = metadados['objeto']
            dados['descricao'] = limpar_texto(corrigir_encoding(descricao_original), max_length=500)

        if metadados.get('data_inicio_propostas'):
            dados['data_leilao'] = metadados['data_inicio_propostas'].split('T')[0]
        elif metadados.get('dataAberturaPropostas'):
            dados['data_leilao'] = metadados['dataAberturaPropostas'].split('T')[0]
        elif descricao_original:
            data_extraida = extrair_data_de_texto(descricao_original)
            if data_extraida:
                dados['data_leilao'] = data_extraida

        urls_todas = []
        texto_completo = ""

        if descricao_original:
            urls_todas.extend(extrair_urls_de_texto(descricao_original))
            texto_completo += descricao_original + " "

        if metadados.get('objeto'):
            urls_todas.extend(extrair_urls_de_texto(metadados['objeto']))
            texto_completo += metadados['objeto'] + " "

        if metadados.get('informacoes_complementares'):
            info_str = str(metadados['informacoes_complementares'])
            urls_todas.extend(extrair_urls_de_texto(info_str))
            texto_completo += info_str + " "

        for file_info in metadados.get('files_meta', []):
            if file_info.get('titulo'):
                urls_todas.extend(extrair_urls_de_texto(file_info['titulo']))

        link_leiloeiro = encontrar_link_leiloeiro(urls_todas, texto_completo)
        if link_leiloeiro:
            dados['link_leiloeiro'] = link_leiloeiro

        if metadados.get('link_pncp'):
            dados['link_pncp'] = metadados['link_pncp']

        if metadados.get('orgao_nome'):
            dados['orgao'] = metadados['orgao_nome']

        if metadados.get('uf'):
            dados['uf'] = metadados['uf']

        if metadados.get('municipio'):
            dados['cidade'] = metadados['municipio']

    except Exception as e:
        print(f"[AVISO] Erro ao ler JSON: {e}")

    return dados


def extrair_de_excel(pasta_dados: Path) -> Dict[str, str]:
    """Prioridade 2: Extrai dados de arquivos Excel."""
    dados = {}

    arquivos_excel = (
        list(pasta_dados.glob("*.xlsx")) +
        list(pasta_dados.glob("*.xls")) +
        list(pasta_dados.glob("*.csv"))
    )

    if not arquivos_excel:
        return dados

    for excel_path in arquivos_excel:
        try:
            if excel_path.suffix == '.csv':
                df = pd.read_csv(excel_path, encoding='utf-8', nrows=50)
            else:
                df = pd.read_excel(excel_path, nrows=50)

            if df.empty:
                continue

            for col in df.columns:
                col_lower = str(col).lower()

                if 'edital' in col_lower and 'n_edital' not in dados:
                    valores = df[col].dropna().astype(str).tolist()
                    if valores:
                        dados['n_edital'] = valores[0]

                if any(kw in col_lower for kw in ['data', 'leilao', 'abertura']):
                    if 'data_leilao' not in dados:
                        valores = df[col].dropna().astype(str).tolist()
                        if valores:
                            data = extrair_data_de_texto(valores[0])
                            if data:
                                dados['data_leilao'] = data

                if any(kw in col_lower for kw in ['desc', 'objeto', 'bem']):
                    if 'descricao' not in dados:
                        valores = df[col].dropna().astype(str).tolist()
                        if valores:
                            dados['descricao'] = ' | '.join(valores[:3])

        except Exception:
            pass

    return dados


def extrair_de_docx(pasta_dados: Path) -> Dict[str, str]:
    """Prioridade 2.5: Extrai dados de arquivos DOCX."""
    dados = {}

    if not DOCX_AVAILABLE:
        return dados

    arquivos_docx = list(pasta_dados.glob("*.docx")) + list(pasta_dados.glob("*.doc"))

    for docx_path in arquivos_docx:
        texto = extrair_texto_docx(docx_path)
        if not texto:
            continue

        if 'n_edital' not in dados:
            match = REGEX_N_EDITAL.search(texto)
            if match:
                dados['n_edital'] = match.group(1).replace(' ', '')

        if 'data_leilao' not in dados:
            data = extrair_data_de_texto(texto)
            if data:
                dados['data_leilao'] = data

        if 'link_leiloeiro' not in dados:
            urls = extrair_urls_de_texto(texto)
            link = encontrar_link_leiloeiro(urls, texto)
            if link:
                dados['link_leiloeiro'] = link

    return dados


def extrair_de_path(pasta_edital: Path) -> Dict[str, str]:
    """Prioridade 3: Extrai dados do nome da pasta principal."""
    dados = {}
    partes = pasta_edital.name.split('_')

    if len(partes) >= 3:
        dados['uf'] = partes[0].upper()
        dados['cidade'] = partes[1].replace('-', ' ').title()
        dados['orgao'] = ' '.join(partes[2:]).replace('-', ' ').title()
    elif len(partes) >= 2:
        dados['uf'] = partes[0].upper()
        dados['cidade'] = partes[1].replace('-', ' ').title()

    return dados


def extrair_de_pdf(pasta_dados: Path) -> Dict[str, str]:
    """Prioridade 4: Extrai dados do PDF."""
    dados = {}

    arquivos_pdf = list(pasta_dados.glob("*.pdf"))
    if not arquivos_pdf:
        return dados

    pdf_path = arquivos_pdf[0]

    try:
        with pdfplumber.open(pdf_path) as pdf:
            texto_completo = ""
            for page in pdf.pages:
                texto_completo += page.extract_text() or ""

            if not texto_completo.strip():
                return dados

            match = REGEX_N_EDITAL.search(texto_completo)
            if match:
                dados['n_edital'] = match.group(1).replace(' ', '')

            data_extraida = extrair_data_de_texto(texto_completo)
            if data_extraida:
                dados['data_leilao'] = data_extraida

            match_titulo = re.search(
                r"(?i)(?:objeto|finalidade).*?[:;]\s*(.{10,200}?)(?:\.|;|\n\n)",
                texto_completo,
                re.DOTALL
            )
            if match_titulo:
                dados['titulo'] = match_titulo.group(1).strip()[:200]

            urls = extrair_urls_de_texto(texto_completo)
            urls_ctx = extrair_urls_contextuais(texto_completo)
            urls.extend(urls_ctx)

            link_leiloeiro = encontrar_link_leiloeiro(urls, texto_completo)
            if link_leiloeiro:
                dados['link_leiloeiro'] = link_leiloeiro

    except Exception as e:
        print(f"[AVISO] Erro ao processar PDF: {e}")

    return dados


def extrair_urls_de_todos_arquivos(pasta_dados: Path) -> tuple:
    """Extrai URLs de todos os tipos de arquivo e retorna texto combinado."""
    urls_todas = []
    texto_completo = ""

    urls_docx = extrair_urls_de_docx(pasta_dados)
    urls_todas.extend(urls_docx)

    urls_xlsx = extrair_urls_de_xlsx(pasta_dados)
    urls_todas.extend(urls_xlsx)

    urls_zip = extrair_urls_de_zip(pasta_dados)
    urls_todas.extend(urls_zip)

    for docx_path in list(pasta_dados.glob("*.docx")) + list(pasta_dados.glob("*.doc")):
        texto_completo += extrair_texto_docx(docx_path) + " "

    for pdf_path in pasta_dados.glob("*.pdf"):
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    texto_completo += (page.extract_text() or "") + " "
        except Exception:
            pass

    return urls_todas, texto_completo


def gerar_tags(orgao: str, uf: str) -> str:
    """Gera tags baseadas no orgao."""
    if not orgao or orgao == 'N/D':
        return 'veiculos_gerais'

    orgao_lower = orgao.lower()

    for palavra_chave, tag in TAGS_POR_ORGAO.items():
        if palavra_chave in orgao_lower:
            return tag

    return 'veiculos_gerais'


# ============================================================================
# PROCESSAMENTO PRINCIPAL
# ============================================================================

def processar_edital(pasta_edital: Path) -> Optional[Dict[str, str]]:
    """Processa um edital usando cascata de fontes."""
    print(f"[INFO] Processando: {pasta_edital.name}")

    pasta_dados = encontrar_pasta_dados(pasta_edital)

    if not pasta_dados:
        print("  [ERRO] Pasta de dados nao encontrada")
        return None

    if pasta_dados != pasta_edital:
        print(f"  [INFO] Pasta de dados: {pasta_dados.name}")

    dados_finais = {
        'n_edital': 'N/D',
        'data_leilao': 'N/D',
        'titulo': 'N/D',
        'descricao': 'N/D',
        'orgao': 'N/D',
        'uf': 'N/D',
        'cidade': 'N/D',
        'tags': 'N/D',
        'link_leiloeiro': 'N/D',
        'link_pncp': 'N/D',
        'arquivo_origem': pasta_edital.name
    }

    fontes = [
        ('JSON', extrair_de_metadados_json(pasta_dados)),
        ('Excel', extrair_de_excel(pasta_dados)),
        ('DOCX', extrair_de_docx(pasta_dados)),
        ('Path', extrair_de_path(pasta_edital)),
        ('PDF', extrair_de_pdf(pasta_dados))
    ]

    campos_encontrados = []
    for nome_fonte, dados_fonte in fontes:
        for campo, valor in dados_fonte.items():
            if campo in dados_finais:
                if dados_finais[campo] == 'N/D' or not dados_finais[campo]:
                    if valor and valor != 'N/D':
                        dados_finais[campo] = valor
                        campos_encontrados.append(f"{campo}({nome_fonte})")

    if dados_finais['link_leiloeiro'] == 'N/D':
        urls_todos, texto_todos = extrair_urls_de_todos_arquivos(pasta_dados)
        link_encontrado = encontrar_link_leiloeiro(urls_todos, texto_todos)
        if link_encontrado:
            dados_finais['link_leiloeiro'] = link_encontrado
            campos_encontrados.append("link_leiloeiro(V9_DEEP)")

    if campos_encontrados:
        preview = campos_encontrados[:5]
        suffix = '...' if len(campos_encontrados) > 5 else ''
        print(f"  [OK] Extraido: {', '.join(preview)}{suffix}")
    else:
        print("  [AVISO] Nenhum dado extraido")

    dados_finais['tags'] = gerar_tags(dados_finais['orgao'], dados_finais['uf'])

    return dados_finais


def varrer_editais() -> List[Dict[str, str]]:
    """Varre todos os editais."""
    if not PASTA_EDITAIS.exists():
        print(f"[ERRO] Diretorio {PASTA_EDITAIS} nao encontrado")
        return []

    print(f"[INFO] Varrendo editais em: {PASTA_EDITAIS}")

    resultados = []
    pastas_editais = [d for d in PASTA_EDITAIS.iterdir() if d.is_dir()]

    print(f"[INFO] Total de pastas encontradas: {len(pastas_editais)}\n")

    for i, pasta in enumerate(pastas_editais, 1):
        try:
            print(f"[{i}/{len(pastas_editais)}] ", end='')
            dados = processar_edital(pasta)
            if dados:
                resultados.append(dados)
        except Exception as e:
            print(f"[ERRO] Falha ao processar {pasta.name}: {e}")

    return resultados


def salvar_csv(dados: List[Dict[str, str]]):
    """Salva resultados em CSV."""
    if not dados:
        print("[AVISO] Nenhum dado para salvar")
        return

    df = pd.DataFrame(dados)
    df.to_csv(CSV_OUTPUT, index=False, encoding='utf-8-sig')

    print(f"\n[OK] Arquivo gerado: {CSV_OUTPUT}")
    print(f"[OK] Total de editais processados: {len(dados)}")

    print("\n" + "=" * 60)
    print("ESTATISTICAS DE PREENCHIMENTO - V9")
    print("=" * 60)

    for col in df.columns:
        if col == 'arquivo_origem':
            continue

        total = len(df)
        nao_vazio = df[col].notna() & (df[col] != '') & (df[col] != 'N/D')
        preenchidos = nao_vazio.sum()
        taxa = (preenchidos / total * 100) if total > 0 else 0

        status = "[OK]" if taxa == 100 else "[PARCIAL]" if taxa >= 50 else "[BAIXO]"
        barra = "#" * int(taxa / 5)
        print(f"{status} {col:20s} {preenchidos:3d}/{total:3d} ({taxa:5.1f}%) {barra}")


# ============================================================================
# MAIN
# ============================================================================

def main():
    print("\n" + "=" * 60)
    print("ACHE SUCATAS DaaS - AUDITOR V9")
    print("=" * 60)
    print("\nMelhorias V9:")
    print("- Fallback hierarquico para link_leiloeiro")
    print("- Regex de URLs expandido")
    print("- Keywords leiloeiro: +10 plataformas")
    print("- Extracao contextual (plataforma, acesse, portal)")
    print("\nCascata: JSON -> Excel -> DOCX -> Path -> PDF")

    if not DOCX_AVAILABLE:
        print("\n[AVISO] python-docx nao instalado!")
        print("Instale com: pip install python-docx")

    print("=" * 60 + "\n")

    resultados = varrer_editais()
    salvar_csv(resultados)

    print("\n" + "=" * 60)
    print("PROCESSAMENTO CONCLUIDO - V9")
    print("=" * 60 + "\n")


if __name__ == "__main__":
    main()
