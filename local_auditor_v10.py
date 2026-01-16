#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
ACHE SUCATAS DaaS - AUDITOR V10
Extrai dados de editais usando cascata: JSON -> Excel -> DOCX -> Path -> PDF

V10:
- Descoberta correta de editais: Município -> subpastas de editais -> arquivos
- Regex de URLs com protocolo opcional (captura www. sem http/https)
- Suporte a .net.br nos padrões de URL
- PDF: itera por todos os PDFs para extração de campos
"""

from __future__ import annotations

import json
import re
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple

import pandas as pd
import pdfplumber

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[AVISO] python-docx nao instalado. Instale com: pip install python-docx")


PASTA_RAIZ = Path(__file__).parent
PASTA_EDITAIS = PASTA_RAIZ / "ACHE_SUCATAS_DB"
CSV_OUTPUT = PASTA_RAIZ / "analise_editais_v10.csv"


REGEX_N_EDITAL = re.compile(
    r"(?i)(?:edital|processo|leilao|pregao).*?(\d{1,5}\s*/\s*20\d{2})",
    re.DOTALL,
)

REGEX_DATA = re.compile(r"\b(\d{1,2}[/\-]\d{1,2}[/\-]20\d{2})\b")

REGEX_DATA_CONTEXTUAL = re.compile(
    r"(?i)(?:data|abertura|sessao|leilao|pregao|realizacao|"
    r"desfazimento|alienacao|hasta|arrematacao).*?"
    r"(\d{1,2}[/\-]\d{1,2}[/\-]20\d{2})",
    re.DOTALL,
)

SUPPORTED_TLDS = r"(?:\.com\.br|\.com|\.org\.br|\.gov\.br|\.net\.br|\.net)"

REGEX_URL = re.compile(
    rf"(?:https?://)?(?:www\.)?[\w\-\.]+{SUPPORTED_TLDS}"
    r"(?:/[\w\-\./?%&=~#]*)?",
    re.IGNORECASE,
)

REGEX_URL_QUEBRADA = re.compile(
    rf"(?:https?://)?(?:www\.)?[\w\-\.]+\s*[\n\r]\s*{SUPPORTED_TLDS}"
    r"(?:/[\w\-\./?%&=~#]*)?",
    re.IGNORECASE,
)

REGEX_URL_ENCURTADA = re.compile(
    r"(?:bit\.ly|goo\.gl|tinyurl\.com|t\.co)/[\w\-]+",
    re.IGNORECASE,
)

REGEX_PLATAFORMA_CONTEXTUAL = re.compile(
    rf"(?i)(?:plataforma|acesse|site|portal|sistema|endereco\s*eletronico|"
    rf"disponivel\s*em|link|url)[:\s]*"
    rf"((?:https?://)?(?:www\.)?[\w\-\.]+{SUPPORTED_TLDS}"
    r"(?:/[\w\-\./?%&=~#]*)?)",
    re.IGNORECASE,
)

TAGS_POR_ORGAO = {
    "prefeitura": "veiculos_municipais",
    "camara": "veiculos_municipais",
    "detran": "veiculos_detran",
    "policia": "veiculos_policiais",
    "tribunal": "veiculos_judiciarios",
    "tj": "veiculos_judiciarios",
    "der": "veiculos_estaduais",
    "departamento de estradas": "veiculos_estaduais",
    "receita federal": "bens_apreendidos_federal",
    "secretaria da receita": "bens_apreendidos_federal",
    "srfb": "bens_apreendidos_federal",
    "patio": "veiculos_detran",
    "custodia": "veiculos_detran",
    "apreendidos": "veiculos_detran",
    "pm": "veiculos_policiais",
    "prf": "veiculos_policiais",
    "policia rodoviaria": "veiculos_policiais",
    "bombeiro": "veiculos_bombeiros",
    "samu": "veiculos_saude",
    "saude": "veiculos_saude",
}

KEYWORDS_LEILOEIRO = [
    "leiloeiro",
    "leilao",
    "lance",
    "arrematacao",
    "superbid",
    "sodresantoro",
    "zukerman",
    "joaoemilio",
    "leiloesfreire",
    "frfreiloes",
    "leilaobrasil",
    "megaleiloes",
    "sold",
    "leilomaster",
    "vipleiloes",
    "leiloesja",
    "portalleiloes",
    "lanceja",
    "leiloar",
    "leiloado",
    "hastaleiloes",
    "cleiloes",
    "bfreiloes",
    "alfreidoleiloes",
    "rioleiloes",
    "leilonet",
    "leilocentro",
]


@dataclass(frozen=True)
class ResultadoEdital:
    n_edital: str = "N/D"
    data_leilao: str = "N/D"
    titulo: str = "N/D"
    descricao: str = "N/D"
    orgao: str = "N/D"
    uf: str = "N/D"
    cidade: str = "N/D"
    tags: str = "N/D"
    link_leiloeiro: str = "N/D"
    link_pncp: str = "N/D"
    arquivo_origem: str = "N/D"

    def as_dict(self) -> Dict[str, str]:
        return {
            "n_edital": self.n_edital,
            "data_leilao": self.data_leilao,
            "titulo": self.titulo,
            "descricao": self.descricao,
            "orgao": self.orgao,
            "uf": self.uf,
            "cidade": self.cidade,
            "tags": self.tags,
            "link_leiloeiro": self.link_leiloeiro,
            "link_pncp": self.link_pncp,
            "arquivo_origem": self.arquivo_origem,
        }


def corrigir_encoding(texto: str) -> str:
    if not texto or texto == "N/D":
        return texto
    try:
        return texto.encode("latin-1").decode("utf-8")
    except (UnicodeDecodeError, UnicodeEncodeError):
        return texto


def limpar_texto(texto: str, max_length: int = 500) -> str:
    if not texto or texto == "N/D":
        return texto
    texto = re.sub(r"\n{3,}", "\n\n", texto)
    texto = re.sub(r" {2,}", " ", texto)
    if len(texto) > max_length:
        texto = texto[:max_length] + "..."
    return texto.strip()


def extrair_data_de_texto(texto: str) -> Optional[str]:
    if not texto:
        return None

    matches = REGEX_DATA_CONTEXTUAL.findall(texto)
    if not matches:
        matches = REGEX_DATA.findall(texto)

    for data_str in matches:
        try:
            data_obj = datetime.strptime(data_str.replace("-", "/"), "%d/%m/%Y")
            if data_obj.year >= 2024:
                return data_str
        except ValueError:
            continue

    return None


def normalizar_url(url: str) -> str:
    url = url.strip()
    url = re.sub(r"\s+", "", url)
    url = url.rstrip(".,;:)>")
    url = url.rstrip("\"'")
    if not url:
        return ""
    if not url.lower().startswith(("http://", "https://")):
        url = "https://" + url
    return url


def extrair_urls_de_texto(texto: str) -> List[str]:
    if not texto:
        return []

    urls: List[str] = []

    urls.extend(REGEX_URL.findall(texto))

    for raw in REGEX_URL_QUEBRADA.findall(texto):
        urls.append(raw)

    for raw in REGEX_URL_ENCURTADA.findall(texto):
        urls.append(raw)

    urls_limpas: List[str] = []
    for raw in urls:
        url = normalizar_url(raw)
        if url and len(url) > 10:
            urls_limpas.append(url)

    return list(dict.fromkeys(urls_limpas))


def extrair_urls_contextuais(texto: str) -> List[str]:
    if not texto:
        return []

    urls: List[str] = []
    for match in REGEX_PLATAFORMA_CONTEXTUAL.findall(texto):
        url = normalizar_url(match)
        if url and len(url) > 10:
            urls.append(url)

    return list(dict.fromkeys(urls))


def encontrar_link_leiloeiro(urls: Sequence[str], texto_completo: str = "") -> Optional[str]:
    urls_unicas = list(dict.fromkeys([u for u in urls if u]))

    for url in urls_unicas:
        url_lower = url.lower()
        if any(keyword in url_lower for keyword in KEYWORDS_LEILOEIRO):
            return url

    if texto_completo:
        urls_ctx = extrair_urls_contextuais(texto_completo)
        for url in urls_ctx:
            url_lower = url.lower()
            if "gov" in url_lower or "pncp" in url_lower:
                continue
            if any(keyword in url_lower for keyword in KEYWORDS_LEILOEIRO):
                return url

        for url in urls_ctx:
            url_lower = url.lower()
            if "gov" in url_lower or "pncp" in url_lower:
                continue
            if ".com.br" in url_lower or ".com" in url_lower or ".net.br" in url_lower or ".net" in url_lower:
                return url

    for url in urls_unicas:
        url_lower = url.lower()
        if "gov" in url_lower or "pncp" in url_lower:
            continue
        if any(tld in url_lower for tld in [".com.br", ".com", ".net.br", ".net", ".org.br"]):
            return url

    return None


def gerar_tags(orgao: str) -> str:
    if not orgao or orgao == "N/D":
        return "veiculos_gerais"

    orgao_lower = orgao.lower()
    for palavra_chave, tag in TAGS_POR_ORGAO.items():
        if palavra_chave in orgao_lower:
            return tag

    return "veiculos_gerais"


def extrair_texto_docx(arquivo_docx: Path) -> str:
    if not DOCX_AVAILABLE:
        return ""
    try:
        doc = DocxDocument(str(arquivo_docx))
        partes: List[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                partes.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        partes.append(cell.text)
        return "\n".join(partes)
    except Exception:
        return ""


def extrair_urls_de_docx(pasta: Path) -> List[str]:
    urls: List[str] = []
    for docx_path in list(pasta.glob("*.docx")) + list(pasta.glob("*.doc")):
        texto = extrair_texto_docx(docx_path)
        if not texto:
            continue
        urls.extend(extrair_urls_de_texto(texto))
        urls.extend(extrair_urls_contextuais(texto))
    return list(dict.fromkeys(urls))


def extrair_urls_de_xlsx(pasta: Path) -> List[str]:
    urls: List[str] = []
    arquivos = list(pasta.glob("*.xlsx")) + list(pasta.glob("*.xls")) + list(pasta.glob("*.xlsm"))
    for excel_path in arquivos:
        try:
            xl = pd.ExcelFile(excel_path)
            for sheet_name in xl.sheet_names:
                df = pd.read_excel(xl, sheet_name=sheet_name, dtype=str)
                for col in df.columns:
                    for valor in df[col].dropna():
                        valor_str = str(valor)
                        if "http" in valor_str.lower() or ".com" in valor_str.lower() or ".net" in valor_str.lower():
                            urls.extend(extrair_urls_de_texto(valor_str))
        except Exception:
            continue
    return list(dict.fromkeys(urls))


def extrair_urls_de_zip(pasta: Path) -> List[str]:
    urls: List[str] = []
    for zip_path in pasta.glob("*.zip"):
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir)
                with zipfile.ZipFile(zip_path, "r") as zf:
                    zf.extractall(temp_path)

                urls.extend(extrair_urls_de_docx(temp_path))
                urls.extend(extrair_urls_de_xlsx(temp_path))

                for pdf_path in temp_path.glob("**/*.pdf"):
                    try:
                        texto = extrair_texto_pdfs([pdf_path])
                        urls.extend(extrair_urls_de_texto(texto))
                        urls.extend(extrair_urls_contextuais(texto))
                    except Exception:
                        continue
        except Exception:
            continue
    return list(dict.fromkeys(urls))


def extrair_texto_pdfs(pdfs: Sequence[Path]) -> str:
    partes: List[str] = []
    for pdf_path in pdfs:
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    partes.append(page.extract_text() or "")
        except Exception:
            continue
    return " ".join([p for p in partes if p]).strip()


def extrair_de_metadados_json(pasta: Path) -> Dict[str, str]:
    dados: Dict[str, str] = {}
    json_file = pasta / "metadados_pncp.json"
    if not json_file.exists():
        return dados

    try:
        with open(json_file, "r", encoding="utf-8") as f:
            metadados = json.load(f)

        titulo_original = metadados.get("titulo") or ""
        if titulo_original:
            titulo_corrigido = corrigir_encoding(titulo_original)
            dados["titulo"] = titulo_corrigido
            match = REGEX_N_EDITAL.search(titulo_corrigido)
            if match:
                dados["n_edital"] = match.group(1).replace(" ", "")

        descricao_original = ""
        if metadados.get("descricao"):
            descricao_original = str(metadados["descricao"])
            dados["descricao"] = limpar_texto(corrigir_encoding(descricao_original), max_length=500)
        elif metadados.get("objeto"):
            descricao_original = str(metadados["objeto"])
            dados["descricao"] = limpar_texto(corrigir_encoding(descricao_original), max_length=500)

        if metadados.get("data_inicio_propostas"):
            dados["data_leilao"] = str(metadados["data_inicio_propostas"]).split("T")[0]
        elif metadados.get("dataAberturaPropostas"):
            dados["data_leilao"] = str(metadados["dataAberturaPropostas"]).split("T")[0]
        elif descricao_original:
            data_extraida = extrair_data_de_texto(descricao_original)
            if data_extraida:
                dados["data_leilao"] = data_extraida

        urls_todas: List[str] = []
        texto_completo = ""

        if descricao_original:
            urls_todas.extend(extrair_urls_de_texto(descricao_original))
            texto_completo += descricao_original + " "

        objeto = str(metadados.get("objeto") or "")
        if objeto:
            urls_todas.extend(extrair_urls_de_texto(objeto))
            texto_completo += objeto + " "

        info = metadados.get("informacoes_complementares")
        if info is not None:
            info_str = str(info)
            urls_todas.extend(extrair_urls_de_texto(info_str))
            texto_completo += info_str + " "

        for file_info in metadados.get("files_meta", []) or []:
            titulo_file = str(file_info.get("titulo") or "")
            if titulo_file:
                urls_todas.extend(extrair_urls_de_texto(titulo_file))

        link_leiloeiro = encontrar_link_leiloeiro(urls_todas, texto_completo)
        if link_leiloeiro:
            dados["link_leiloeiro"] = link_leiloeiro

        link_pncp = metadados.get("link_pncp")
        if link_pncp:
            dados["link_pncp"] = str(link_pncp)

        orgao = metadados.get("orgao_nome")
        if orgao:
            dados["orgao"] = str(orgao)

        uf = metadados.get("uf")
        if uf:
            dados["uf"] = str(uf)

        municipio = metadados.get("municipio")
        if municipio:
            dados["cidade"] = str(municipio)

    except Exception as exc:
        print(f"[AVISO] Erro ao ler JSON: {exc}")

    return dados


def extrair_de_excel(pasta: Path) -> Dict[str, str]:
    dados: Dict[str, str] = {}
    arquivos = list(pasta.glob("*.xlsx")) + list(pasta.glob("*.xls")) + list(pasta.glob("*.csv"))
    if not arquivos:
        return dados

    for arquivo in arquivos:
        try:
            if arquivo.suffix.lower() == ".csv":
                df = pd.read_csv(arquivo, encoding="utf-8", nrows=50)
            else:
                df = pd.read_excel(arquivo, nrows=50)
            if df.empty:
                continue

            for col in df.columns:
                col_lower = str(col).lower()

                if "edital" in col_lower and "n_edital" not in dados:
                    valores = df[col].dropna().astype(str).tolist()
                    if valores:
                        dados["n_edital"] = valores[0]

                if any(kw in col_lower for kw in ["data", "leilao", "abertura"]) and "data_leilao" not in dados:
                    valores = df[col].dropna().astype(str).tolist()
                    if valores:
                        data = extrair_data_de_texto(valores[0])
                        if data:
                            dados["data_leilao"] = data

                if any(kw in col_lower for kw in ["desc", "objeto", "bem"]) and "descricao" not in dados:
                    valores = df[col].dropna().astype(str).tolist()
                    if valores:
                        dados["descricao"] = " | ".join(valores[:3])
        except Exception:
            continue

    return dados


def extrair_de_docx(pasta: Path) -> Dict[str, str]:
    dados: Dict[str, str] = {}
    if not DOCX_AVAILABLE:
        return dados

    arquivos = list(pasta.glob("*.docx")) + list(pasta.glob("*.doc"))
    for docx_path in arquivos:
        texto = extrair_texto_docx(docx_path)
        if not texto:
            continue

        if "n_edital" not in dados:
            match = REGEX_N_EDITAL.search(texto)
            if match:
                dados["n_edital"] = match.group(1).replace(" ", "")

        if "data_leilao" not in dados:
            data = extrair_data_de_texto(texto)
            if data:
                dados["data_leilao"] = data

        if "link_leiloeiro" not in dados:
            urls = extrair_urls_de_texto(texto)
            link = encontrar_link_leiloeiro(urls, texto)
            if link:
                dados["link_leiloeiro"] = link

    return dados


def extrair_de_path(pasta_edital: Path) -> Dict[str, str]:
    dados: Dict[str, str] = {}
    partes = pasta_edital.name.split("_")

    if len(partes) >= 3:
        dados["uf"] = partes[0].upper()
        dados["cidade"] = partes[1].replace("-", " ").title()
        dados["orgao"] = " ".join(partes[2:]).replace("-", " ").title()
    elif len(partes) >= 2:
        dados["uf"] = partes[0].upper()
        dados["cidade"] = partes[1].replace("-", " ").title()

    return dados


def extrair_de_pdf(pasta: Path) -> Dict[str, str]:
    dados: Dict[str, str] = {}
    pdfs = sorted(list(pasta.glob("*.pdf")))
    if not pdfs:
        return dados

    texto_completo = extrair_texto_pdfs(pdfs)
    if not texto_completo:
        return dados

    match = REGEX_N_EDITAL.search(texto_completo)
    if match:
        dados["n_edital"] = match.group(1).replace(" ", "")

    data_extraida = extrair_data_de_texto(texto_completo)
    if data_extraida:
        dados["data_leilao"] = data_extraida

    match_titulo = re.search(
        r"(?i)(?:objeto|finalidade).*?[:;]\s*(.{10,200}?)(?:\.|;|\n\n)",
        texto_completo,
        re.DOTALL,
    )
    if match_titulo:
        dados["titulo"] = match_titulo.group(1).strip()[:200]

    urls = extrair_urls_de_texto(texto_completo)
    urls.extend(extrair_urls_contextuais(texto_completo))
    link = encontrar_link_leiloeiro(urls, texto_completo)
    if link:
        dados["link_leiloeiro"] = link

    return dados


def extrair_urls_de_todos_arquivos(pasta: Path) -> Tuple[List[str], str]:
    urls: List[str] = []
    texto: List[str] = []

    urls.extend(extrair_urls_de_docx(pasta))
    urls.extend(extrair_urls_de_xlsx(pasta))
    urls.extend(extrair_urls_de_zip(pasta))

    for docx_path in list(pasta.glob("*.docx")) + list(pasta.glob("*.doc")):
        doc_text = extrair_texto_docx(docx_path)
        if doc_text:
            texto.append(doc_text)

    pdfs = sorted(list(pasta.glob("*.pdf")))
    pdf_text = extrair_texto_pdfs(pdfs)
    if pdf_text:
        texto.append(pdf_text)

    texto_completo = " ".join(texto).strip()
    return list(dict.fromkeys(urls)), texto_completo


def pasta_parece_edital(pasta: Path) -> bool:
    if not pasta.is_dir():
        return False
    if (pasta / "metadados_pncp.json").exists():
        return True
    padroes = ["*.pdf", "*.docx", "*.doc", "*.xlsx", "*.xls", "*.xlsm", "*.csv", "*.zip"]
    return any(pasta.glob(p) for p in padroes)


def listar_pastas_editais(pasta_raiz: Path) -> List[Path]:
    if not pasta_raiz.exists():
        return []

    editais: List[Path] = []
    municipios = [d for d in pasta_raiz.iterdir() if d.is_dir()]

    for municipio in municipios:
        subdirs = [d for d in municipio.iterdir() if d.is_dir()]
        subdirs_editais = [d for d in subdirs if pasta_parece_edital(d)]

        if subdirs_editais:
            editais.extend(sorted(subdirs_editais))
            continue

        if pasta_parece_edital(municipio):
            editais.append(municipio)

    return editais


def processar_edital(pasta_edital: Path) -> Optional[Dict[str, str]]:
    print(f"[INFO] Processando: {pasta_edital.relative_to(PASTA_EDITAIS)}")

    if not pasta_edital.exists() or not pasta_edital.is_dir():
        print("  [ERRO] Pasta do edital invalida")
        return None

    resultado = ResultadoEdital(arquivo_origem=str(pasta_edital.relative_to(PASTA_EDITAIS)))
    dados_finais = resultado.as_dict()

    fontes: List[Tuple[str, Dict[str, str]]] = [
        ("JSON", extrair_de_metadados_json(pasta_edital)),
        ("Excel", extrair_de_excel(pasta_edital)),
        ("DOCX", extrair_de_docx(pasta_edital)),
        ("Path", extrair_de_path(pasta_edital)),
        ("PDF", extrair_de_pdf(pasta_edital)),
    ]

    campos_encontrados: List[str] = []
    for nome_fonte, dados_fonte in fontes:
        for campo, valor in dados_fonte.items():
            if campo not in dados_finais:
                continue
            if dados_finais[campo] == "N/D" and valor and valor != "N/D":
                dados_finais[campo] = valor
                campos_encontrados.append(f"{campo}({nome_fonte})")

    if dados_finais["link_leiloeiro"] == "N/D":
        urls_todos, texto_todos = extrair_urls_de_todos_arquivos(pasta_edital)
        link_encontrado = encontrar_link_leiloeiro(urls_todos, texto_todos)
        if link_encontrado:
            dados_finais["link_leiloeiro"] = link_encontrado
            campos_encontrados.append("link_leiloeiro(V10_DEEP)")

    dados_finais["tags"] = gerar_tags(dados_finais["orgao"])

    if campos_encontrados:
        preview = campos_encontrados[:5]
        suffix = "..." if len(campos_encontrados) > 5 else ""
        print(f"  [OK] Extraido: {', '.join(preview)}{suffix}")
    else:
        print("  [AVISO] Nenhum dado extraido")

    return dados_finais


def varrer_editais() -> List[Dict[str, str]]:
    if not PASTA_EDITAIS.exists():
        print(f"[ERRO] Diretorio {PASTA_EDITAIS} nao encontrado")
        return []

    print(f"[INFO] Varrendo editais em: {PASTA_EDITAIS}")

    pastas_editais = listar_pastas_editais(PASTA_EDITAIS)
    print(f"[INFO] Total de editais encontrados: {len(pastas_editais)}\n")

    resultados: List[Dict[str, str]] = []
    for i, pasta in enumerate(pastas_editais, 1):
        try:
            print(f"[{i}/{len(pastas_editais)}] ", end="")
            dados = processar_edital(pasta)
            if dados:
                resultados.append(dados)
        except Exception as exc:
            print(f"[ERRO] Falha ao processar {pasta}: {exc}")

    return resultados


def salvar_csv(dados: List[Dict[str, str]]) -> None:
    if not dados:
        print("[AVISO] Nenhum dado para salvar")
        return

    df = pd.DataFrame(dados)
    df.to_csv(CSV_OUTPUT, index=False, encoding="utf-8-sig")

    print(f"\n[OK] Arquivo gerado: {CSV_OUTPUT}")
    print(f"[OK] Total de editais processados: {len(dados)}")

    print("\n" + "=" * 60)
    print("ESTATISTICAS DE PREENCHIMENTO - V10")
    print("=" * 60)

    for col in df.columns:
        if col == "arquivo_origem":
            continue
        total = len(df)
        nao_vazio = df[col].notna() & (df[col] != "") & (df[col] != "N/D")
        preenchidos = int(nao_vazio.sum())
        taxa = (preenchidos / total * 100) if total > 0 else 0

        status = "[OK]" if taxa == 100 else "[PARCIAL]" if taxa >= 50 else "[BAIXO]"
        barra = "#" * int(taxa / 5)
        print(f"{status} {col:20s} {preenchidos:3d}/{total:3d} ({taxa:5.1f}%) {barra}")


def main() -> None:
    print("\n" + "=" * 60)
    print("ACHE SUCATAS DaaS - AUDITOR V10")
    print("=" * 60)
    print("\nCascata: JSON -> Excel -> DOCX -> Path -> PDF")
    print("V10: Municipio -> Editais (subpastas) -> Arquivos")
    print("V10: URLs com/sem protocolo e suporte a .net.br")
    print("V10: PDF itera por todos os PDFs")
    print("=" * 60 + "\n")

    if not DOCX_AVAILABLE:
        print("[AVISO] python-docx nao instalado. DOCX sera ignorado.\n")

    resultados = varrer_editais()
    salvar_csv(resultados)

    print("\n" + "=" * 60)
    print("PROCESSAMENTO CONCLUIDO - V10")
    print("=" * 60 + "\n")


if __name__ == "__main__":
    main()
